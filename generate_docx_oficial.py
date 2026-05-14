#!/usr/bin/env python3
"""
Genera el documento DOCX oficial del Proyecto Intermodular DAM - ReciApp.
Sigue el esquema de documentacion del CFGS DAM.
"""

from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml
import os

# =====================================================================
# CONSTANTES
# =====================================================================
GREEN = RGBColor(0x16, 0xA3, 0x4A)
DARK_GREEN = RGBColor(0x15, 0x80, 0x3D)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)
BLACK = RGBColor(0x00, 0x00, 0x00)
DARK = RGBColor(0x1F, 0x29, 0x37)
RED_BORDER = RGBColor(0xDC, 0x26, 0x26)

fig_counter = 0

# =====================================================================
# HELPERS
# =====================================================================

def set_cell_shading(cell, color_hex):
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color_hex}" w:val="clear"/>')
    cell._tc.get_or_add_tcPr().append(shading)


def fmt(paragraph, align=WD_ALIGN_PARAGRAPH.JUSTIFY, spacing_after=Pt(6),
        spacing_before=Pt(0), line_spacing=1.15):
    """Aplica formato estandar a un parrafo."""
    pf = paragraph.paragraph_format
    pf.alignment = align
    pf.space_after = spacing_after
    pf.space_before = spacing_before
    pf.line_spacing = line_spacing
    return paragraph


def body(doc, text, bold=False, italic=False, size=11):
    """Parrafo de cuerpo justificado."""
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.size = Pt(size)
    run.font.name = 'Calibri'
    run.font.color.rgb = BLACK
    run.bold = bold
    run.italic = italic
    fmt(p)
    return p


def heading1(doc, text):
    """Titulo de seccion principal: MAYUSCULAS, negrita, 2 lineas antes."""
    p = doc.add_paragraph()
    run = p.add_run(text.upper())
    run.font.size = Pt(14)
    run.font.name = 'Calibri'
    run.font.color.rgb = DARK_GREEN
    run.bold = True
    fmt(p, align=WD_ALIGN_PARAGRAPH.LEFT,
        spacing_before=Pt(24), spacing_after=Pt(12))
    return p


def heading2(doc, text):
    """Subtitulo de seccion."""
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.size = Pt(12)
    run.font.name = 'Calibri'
    run.font.color.rgb = DARK_GREEN
    run.bold = True
    fmt(p, align=WD_ALIGN_PARAGRAPH.LEFT,
        spacing_before=Pt(18), spacing_after=Pt(8))
    return p


def heading3(doc, text):
    """Sub-subtitulo."""
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.size = Pt(11)
    run.font.name = 'Calibri'
    run.font.color.rgb = GREEN
    run.bold = True
    fmt(p, align=WD_ALIGN_PARAGRAPH.LEFT,
        spacing_before=Pt(12), spacing_after=Pt(6))
    return p


def bullet(doc, text, level=0):
    """Punto de lista."""
    p = doc.add_paragraph(style='List Bullet')
    p.clear()
    run = p.add_run(text)
    run.font.size = Pt(11)
    run.font.name = 'Calibri'
    run.font.color.rgb = BLACK
    if level > 0:
        p.paragraph_format.left_indent = Cm(1.27 * (level + 1))
    fmt(p, align=WD_ALIGN_PARAGRAPH.JUSTIFY)
    return p


def tabla(doc, headers, rows, col_widths=None):
    """Tabla con cabecera verde."""
    t = doc.add_table(rows=1 + len(rows), cols=len(headers))
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    t.style = 'Table Grid'
    for i, h in enumerate(headers):
        c = t.rows[0].cells[i]
        set_cell_shading(c, "158035")
        p = c.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(h)
        r.font.size = Pt(10)
        r.font.color.rgb = WHITE
        r.bold = True
        r.font.name = 'Calibri'
    for ri, row in enumerate(rows):
        for ci, val in enumerate(row):
            c = t.rows[ri + 1].cells[ci]
            if ri % 2 == 1:
                set_cell_shading(c, "F0FDF4")
            p = c.paragraphs[0]
            r = p.add_run(str(val))
            r.font.size = Pt(10)
            r.font.name = 'Calibri'
            r.font.color.rgb = BLACK
    if col_widths:
        for i, w in enumerate(col_widths):
            for row in t.rows:
                row.cells[i].width = Cm(w)
    doc.add_paragraph()
    return t


def code_block(doc, code_text):
    """Bloque de codigo con fondo gris y borde verde izquierdo."""
    t = doc.add_table(rows=1, cols=1)
    t.alignment = WD_TABLE_ALIGNMENT.LEFT
    cell = t.cell(0, 0)
    set_cell_shading(cell, "F3F4F6")
    tcPr = cell._tc.get_or_add_tcPr()
    tcPr.append(parse_xml(
        f'<w:tcBorders {nsdecls("w")}>'
        f'  <w:top w:val="single" w:sz="4" w:space="0" w:color="D1D5DB"/>'
        f'  <w:left w:val="single" w:sz="12" w:space="0" w:color="16A34A"/>'
        f'  <w:bottom w:val="single" w:sz="4" w:space="0" w:color="D1D5DB"/>'
        f'  <w:right w:val="single" w:sz="4" w:space="0" w:color="D1D5DB"/>'
        f'</w:tcBorders>'
    ))
    p = cell.paragraphs[0]
    r = p.add_run(code_text)
    r.font.size = Pt(9)
    r.font.name = 'Consolas'
    r.font.color.rgb = DARK
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    doc.add_paragraph()


def figura(doc, description):
    """Placeholder de captura con borde rojo discontinuo y numero de figura."""
    global fig_counter
    fig_counter += 1

    t = doc.add_table(rows=1, cols=1)
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = t.cell(0, 0)
    tcPr = cell._tc.get_or_add_tcPr()
    tcPr.append(parse_xml(
        f'<w:tcBorders {nsdecls("w")}>'
        f'  <w:top w:val="dashed" w:sz="12" w:space="0" w:color="DC2626"/>'
        f'  <w:left w:val="dashed" w:sz="12" w:space="0" w:color="DC2626"/>'
        f'  <w:bottom w:val="dashed" w:sz="12" w:space="0" w:color="DC2626"/>'
        f'  <w:right w:val="dashed" w:sz="12" w:space="0" w:color="DC2626"/>'
        f'</w:tcBorders>'
    ))
    p1 = cell.paragraphs[0]
    p1.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r1 = p1.add_run(f"Figura {fig_counter}")
    r1.font.size = Pt(14)
    r1.font.color.rgb = RED_BORDER
    r1.bold = True
    p1.paragraph_format.space_before = Pt(18)

    p2 = cell.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r2 = p2.add_run(description)
    r2.font.size = Pt(11)
    r2.font.color.rgb = RGBColor(0x99, 0x33, 0x33)
    r2.italic = True

    p3 = cell.add_paragraph()
    p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r3 = p3.add_run("\n\n")
    r3.font.size = Pt(20)
    p3.paragraph_format.space_after = Pt(10)

    # Pie de figura
    pf = doc.add_paragraph()
    pf.alignment = WD_ALIGN_PARAGRAPH.CENTER
    rf = pf.add_run(f"Figura {fig_counter}: {description}")
    rf.font.size = Pt(10)
    rf.font.name = 'Calibri'
    rf.font.color.rgb = RGBColor(0x55, 0x55, 0x55)
    rf.italic = True
    pf.paragraph_format.space_after = Pt(12)

    return fig_counter


# =====================================================================
# DOCUMENTO
# =====================================================================
doc = Document()

# Estilo base
style = doc.styles['Normal']
style.font.name = 'Calibri'
style.font.size = Pt(11)
style.font.color.rgb = BLACK
style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
style.paragraph_format.line_spacing = 1.15

# Margenes A4 por defecto
for section in doc.sections:
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin = Cm(3)
    section.right_margin = Cm(3)


# =================================================================
#  PORTADA
# =================================================================
for _ in range(5):
    doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("PROYECTO INTERMODULAR")
r.font.size = Pt(16)
r.font.name = 'Calibri'
r.bold = True
r.font.color.rgb = DARK_GREEN

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("CFGS Desarrollo de Aplicaciones Multiplataforma (DAM)")
r.font.size = Pt(13)
r.font.name = 'Calibri'
r.font.color.rgb = BLACK

for _ in range(2):
    doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("ReciApp")
r.font.size = Pt(36)
r.font.name = 'Calibri'
r.bold = True
r.font.color.rgb = DARK_GREEN

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("Sistema de Gestion de Reciclaje y Reduccion de Emisiones CO2")
r.font.size = Pt(14)
r.font.name = 'Calibri'
r.font.color.rgb = RGBColor(0x6B, 0x72, 0x80)
r.italic = True

for _ in range(4):
    doc.add_paragraph()

info = [
    ("Alumno:", "[Nombre del alumno]"),
    ("Tutor:", "[Nombre del tutor]"),
    ("Centro:", "[Nombre del centro educativo]"),
    ("Curso:", "2025 / 2026"),
    ("Fecha:", "Mayo 2026"),
]
for label, value in info:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r1 = p.add_run(label + "  ")
    r1.bold = True
    r1.font.size = Pt(12)
    r1.font.name = 'Calibri'
    r2 = p.add_run(value)
    r2.font.size = Pt(12)
    r2.font.name = 'Calibri'
    p.paragraph_format.space_after = Pt(2)

doc.add_page_break()


# =================================================================
#  INDICE
# =================================================================
heading1(doc, "Indice de contenidos")

indice = [
    "1. Analisis del problema",
    "    1.1. Introduccion",
    "    1.2. Objetivos principales",
    "    1.3. Funciones y rendimientos deseados",
    "    1.4. Planteamiento y evaluacion de diversas soluciones",
    "    1.5. Justificacion de la solucion elegida",
    "    1.6. Modelado de la solucion",
    "        1.6.1. Recursos humanos",
    "        1.6.2. Recursos hardware",
    "        1.6.3. Recursos software",
    "    1.7. Planificacion temporal",
    "2. Diseno e implementacion del proyecto",
    "    2.1. Plan de empresa",
    "    2.2. Arquitectura de la aplicacion",
    "    2.3. Roles de usuario",
    "    2.4. Mapa de navegacion",
    "    2.5. Base de datos",
    "    2.6. Implementacion por modulos",
    "3. Fase de pruebas",
    "4. Documentacion de la aplicacion",
    "    4.1. Introduccion a la aplicacion",
    "    4.2. Manual de instalacion",
    "    4.3. Manual de usuario",
    "    4.4. Manual de administracion",
    "5. Conclusiones finales",
    "6. Bibliografia",
]
for item in indice:
    p = doc.add_paragraph()
    r = p.add_run(item)
    r.font.size = Pt(11)
    r.font.name = 'Calibri'
    p.paragraph_format.space_after = Pt(2)

doc.add_page_break()


# =================================================================
#  1. ANALISIS DEL PROBLEMA
# =================================================================
heading1(doc, "1. Analisis del problema")

# ---- 1.1 Introduccion ----
heading2(doc, "1.1. Introduccion")

body(doc,
    "ReciApp es un sistema integral de gestion de reciclaje desarrollado "
    "como proyecto intermodular del CFGS de Desarrollo de Aplicaciones "
    "Multiplataforma (DAM). El proyecto aborda la necesidad de fomentar "
    "el reciclaje entre la poblacion mediante la gamificacion y el "
    "seguimiento del impacto ambiental individual, medido en kilogramos "
    "de CO2 reducidos.")

body(doc,
    "El sistema se compone de cuatro aplicaciones interconectadas: una API "
    "REST centralizada desarrollada con Spring Boot, una aplicacion movil "
    "multiplataforma construida con React Native y Expo, una aplicacion de "
    "escritorio con interfaz grafica desarrollada en JavaFX, y un script de "
    "terminal en Python pensado para puntos de reciclaje fisicos. Todas las "
    "aplicaciones se comunican con la API REST, que a su vez gestiona la "
    "persistencia de datos en una base de datos MySQL.")

body(doc,
    "El funcionamiento del sistema es el siguiente: cada producto reciclable "
    "registrado en el catalogo tiene asociado un valor de emisiones de CO2 "
    "reducibles. Cuando un usuario recicla un producto (identificado mediante "
    "su codigo de barras), dicho valor se acumula en su perfil personal, "
    "permitiendole visualizar su contribucion individual a la reduccion de "
    "emisiones de gases de efecto invernadero.")


# ---- 1.2 Objetivos principales ----
heading2(doc, "1.2. Objetivos principales")

body(doc,
    "Se han establecido los siguientes objetivos principales para el desarrollo "
    "del proyecto:")

bullet(doc, "Desarrollar un sistema multiplataforma completo que permita registrar, consultar y gestionar acciones de reciclaje de productos.")
bullet(doc, "Implementar un mecanismo de seguimiento de emisiones de CO2 reducidas por cada usuario, fomentando la concienciacion medioambiental.")
bullet(doc, "Disenar una arquitectura cliente-servidor robusta con una API REST central que permita la comunicacion entre todos los clientes.")
bullet(doc, "Crear una aplicacion movil accesible y facil de usar para consultar productos y visualizar el impacto ambiental personal.")
bullet(doc, "Desarrollar una aplicacion de escritorio con funcionalidades de administracion completas (CRUD de usuarios, productos y transacciones).")
bullet(doc, "Implementar un terminal de reciclaje para puntos fisicos que utilice codigos de barras y un token personal (TAP) para identificar al usuario.")
bullet(doc, "Garantizar la seguridad del sistema mediante autenticacion JWT y cifrado de contrasenas con BCrypt.")
bullet(doc, "Desplegar el sistema en un servidor en la nube (AWS EC2) para demostrar su viabilidad en un entorno real.")


# ---- 1.3 Funciones y rendimientos deseados ----
heading2(doc, "1.3. Funciones y rendimientos deseados")

body(doc, "El sistema implementado ofrece los siguientes servicios:")

heading3(doc, "Gestion de usuarios")
bullet(doc, "Registro y autenticacion con credenciales (nombre de usuario y contrasena).")
bullet(doc, "Gestion de perfil personal: modificacion de nombre, cambio de contrasena y eliminacion de cuenta.")
bullet(doc, "Sistema de roles diferenciados: cliente (usuario estandar) y administrador.")
bullet(doc, "Generacion de TAP (Token de Autenticacion Personal): numero unico de 6 digitos para identificacion en terminales fisicos.")

heading3(doc, "Gestion de productos")
bullet(doc, "Catalogo de productos reciclables con informacion detallada: nombre, material, codigo de barras, emisiones reducibles e imagen.")
bullet(doc, "Busqueda por nombre o codigo de barras y filtrado por tipo de material.")
bullet(doc, "Operaciones CRUD completas para administradores.")

heading3(doc, "Registro de reciclaje")
bullet(doc, "Registro de acciones de reciclaje asociando un producto a un usuario con fecha y hora exactas.")
bullet(doc, "Acumulacion automatica de emisiones de CO2 reducidas en el perfil del usuario.")
bullet(doc, "Historial completo de reciclajes con estadisticas de impacto ambiental.")

heading3(doc, "Funcionalidades adicionales")
bullet(doc, "Modo offline/invitado en la aplicacion movil con datos de ejemplo almacenados localmente.")
bullet(doc, "Graficos de emisiones por usuario en la aplicacion de escritorio.")
bullet(doc, "Interfaz de terminal con colores y formato visual para puntos de reciclaje.")

heading3(doc, "Rendimiento esperado")
bullet(doc, "Tiempo de respuesta de la API inferior a 500 ms para operaciones estandar.")
bullet(doc, "Soporte de multiples conexiones simultaneas desde los distintos clientes.")
bullet(doc, "Disponibilidad 24/7 gracias al despliegue en servidor cloud.")


# ---- 1.4 Planteamiento y evaluacion de diversas soluciones ----
heading2(doc, "1.4. Planteamiento y evaluacion de diversas soluciones")

body(doc,
    "Durante la fase de planificacion se evaluaron distintas aproximaciones "
    "para cada componente del sistema:")

heading3(doc, "Backend / API")
body(doc,
    "Se consideraron tres alternativas principales para el backend: "
    "Node.js con Express, Django REST Framework (Python) y Spring Boot (Java). "
    "Node.js ofrecia simplicidad pero menor tipado y robustez. Django era una "
    "opcion viable con buena documentacion pero menor familiaridad por parte "
    "del desarrollador. Spring Boot, por su parte, proporcionaba un ecosistema "
    "maduro con Spring Security, Spring Data JPA y una integracion natural con "
    "las aplicaciones Java del proyecto.")

heading3(doc, "Aplicacion movil")
body(doc,
    "Se evaluaron Flutter (Dart), desarrollo nativo (Kotlin/Swift) y React "
    "Native (TypeScript). Flutter ofrecia buen rendimiento pero requeria "
    "aprender Dart. El desarrollo nativo implicaba mantener dos bases de "
    "codigo. React Native con Expo permitia desarrollo multiplataforma con "
    "un lenguaje ampliamente conocido (TypeScript) y un ecosistema de "
    "herramientas que agiliza el desarrollo.")

heading3(doc, "Aplicacion de escritorio")
body(doc,
    "Se compararon Electron (JavaScript), .NET WPF (C#) y JavaFX (Java). "
    "Electron generaba aplicaciones pesadas y duplicaba tecnologias con la "
    "app movil. WPF limitaba la plataforma a Windows. JavaFX era "
    "multiplataforma, se integraba con el backend Java y cumplia con los "
    "requisitos del modulo de Desarrollo de Interfaces.")

heading3(doc, "Base de datos")
body(doc,
    "Se evaluaron PostgreSQL, MySQL y MongoDB. MongoDB, al ser NoSQL, no "
    "se ajustaba a los requisitos del modulo de Acceso a Datos que requiere "
    "bases de datos relacionales. PostgreSQL y MySQL eran ambas solidas; se "
    "selecciono MySQL por su mayor implantacion en entornos educativos y "
    "compatibilidad con las herramientas del ciclo.")

heading3(doc, "Comunicacion")
body(doc,
    "Se evaluo el uso de HTTPS con certificados autofirmados frente a HTTP "
    "plano. La implementacion de HTTPS con certificados autofirmados resulto "
    "incompatible con Expo Go (React Native), ya que esta plataforma no "
    "permite personalizar la configuracion de seguridad de red de Android "
    "sin eyectar del flujo gestionado. Por esta razon se opto por HTTP para "
    "el entorno de desarrollo, documentando la recomendacion de usar un "
    "certificado de una CA reconocida (Let's Encrypt) en produccion.")


# ---- 1.5 Justificacion de la solucion elegida ----
heading2(doc, "1.5. Justificacion de la solucion elegida")

body(doc,
    "La solucion final elegida se compone de las siguientes tecnologias:")

tabla(doc,
    ["Componente", "Tecnologia elegida", "Justificacion"],
    [
        ["Backend", "Spring Boot 3.4.4 (Java 21)",
         "Ecosistema maduro, Spring Security para JWT, Spring Data JPA, compatibilidad con la app JavaFX."],
        ["App movil", "React Native 0.81.5 + Expo SDK 54",
         "Desarrollo multiplataforma real (Android/iOS), TypeScript, hot reload, Expo Go para desarrollo rapido."],
        ["App escritorio", "JavaFX 25 (Java 24)",
         "Multiplataforma, interfaz grafica rica con FXML/CSS, modulo de Desarrollo de Interfaces."],
        ["Terminal", "Python 3 + requests",
         "Rapidez de desarrollo, ideal para scripts de terminal, colores ANSI nativos."],
        ["Base de datos", "MySQL 8.x (InnoDB)",
         "Relacional, transaccional, amplia compatibilidad, requisito del modulo de Acceso a Datos."],
        ["Comunicacion", "HTTP + JWT",
         "API REST stateless, tokens con expiracion de 7 dias, BCrypt para contrasenas."],
        ["Despliegue", "AWS EC2 (Ubuntu 22.04)",
         "Servidor cloud real, accesible desde todos los clientes, preparado para Docker."],
    ],
    col_widths=[2.8, 4.5, 7]
)

body(doc,
    "Esta combinacion de tecnologias cubre la totalidad de los modulos del "
    "ciclo DAM: Acceso a Datos (Spring Data JPA, MySQL, API REST), "
    "Desarrollo de Interfaces (JavaFX con FXML y CSS), Programacion de "
    "Servicios y Procesos (comunicacion HTTP, JWT, multiples clientes "
    "simultaneos) y Programacion Multimedia y Dispositivos Moviles "
    "(React Native con Expo).")


# ---- 1.6 Modelado de la solucion ----
heading2(doc, "1.6. Modelado de la solucion")

heading3(doc, "1.6.1. Recursos humanos")
body(doc,
    "El proyecto ha sido desarrollado de forma individual por un unico "
    "alumno del CFGS de DAM, asumiendo todos los roles del desarrollo: "
    "analisis, diseno de base de datos, desarrollo backend, desarrollo "
    "frontend (movil y escritorio), scripting (terminal Python), "
    "despliegue en servidor y documentacion.")

heading3(doc, "1.6.2. Recursos hardware")
tabla(doc,
    ["Recurso", "Especificaciones", "Uso"],
    [
        ["Portatil de desarrollo", "ASUS FA506NC / Windows 11", "Desarrollo de todas las aplicaciones, pruebas locales"],
        ["Instancia AWS EC2", "t2.micro / Ubuntu 22.04 LTS", "Alojamiento de la API REST y MySQL en produccion"],
        ["Dispositivo movil Android", "Android 12+", "Pruebas de la app movil con Expo Go"],
    ],
    col_widths=[4, 5, 5.5]
)

heading3(doc, "1.6.3. Recursos software")
tabla(doc,
    ["Software", "Version", "Proposito"],
    [
        ["IntelliJ IDEA / VS Code", "2024/2025", "IDE principal para Java y TypeScript"],
        ["JDK (Temurin)", "21 / 24", "Compilacion del backend y la app de escritorio"],
        ["Node.js", "20 LTS", "Entorno de ejecucion para Expo / React Native"],
        ["Python", "3.10+", "Ejecucion del terminal de reciclaje"],
        ["MySQL Server", "8.x", "Servidor de base de datos"],
        ["MySQL Workbench", "8.x", "Diseno visual del modelo ER"],
        ["Postman", "11.x", "Pruebas de la API REST"],
        ["Git + GitHub", "—", "Control de versiones y repositorio remoto"],
        ["Expo Go", "—", "Ejecucion de la app movil en dispositivo fisico"],
        ["Docker / Docker Compose", "24.x / 2.x", "Contenedorizacion (preparado para produccion)"],
        ["draw.io", "—", "Diagramas de arquitectura y flujos"],
    ],
    col_widths=[4, 2.5, 8]
)


# ---- 1.7 Planificacion temporal ----
heading2(doc, "1.7. Planificacion temporal")

body(doc,
    "La planificacion del proyecto se ha organizado en las siguientes fases:")

tabla(doc,
    ["Fase", "Periodo", "Tareas principales"],
    [
        ["Analisis y diseno", "Febrero 2026 (semanas 1-2)",
         "Definicion de requisitos, diseno del modelo ER, seleccion de tecnologias, creacion del esquema SQL."],
        ["Backend (API REST)", "Febrero - Marzo 2026 (semanas 3-6)",
         "Implementacion de entidades JPA, repositorios, servicios, controladores, seguridad JWT."],
        ["App de escritorio (JavaFX)", "Marzo 2026 (semanas 5-8)",
         "Diseno de interfaces FXML, implementacion de controladores, ApiClient, graficos."],
        ["App movil (React Native)", "Marzo - Abril 2026 (semanas 7-10)",
         "Navegacion con Expo Router, pantallas, stores Zustand, servicio API, modo offline."],
        ["Terminal Python", "Abril 2026 (semana 11)",
         "Script de terminal, integracion con API, interfaz con colores ANSI."],
        ["Despliegue", "Abril 2026 (semana 12)",
         "Configuracion de EC2, despliegue de MySQL y JAR, pruebas de conectividad."],
        ["Pruebas e integracion", "Abril - Mayo 2026 (semanas 12-14)",
         "Pruebas funcionales, correccion de errores, pruebas de integracion entre componentes."],
        ["Documentacion", "Mayo 2026 (semanas 14-16)",
         "Redaccion de la memoria, capturas de pantalla, manuales de usuario e instalacion."],
    ],
    col_widths=[3.5, 4, 7]
)

body(doc,
    "Nota: las fases de backend, escritorio y movil se han solapado "
    "parcialmente, ya que el desarrollo del backend se realizo de forma "
    "incremental a medida que se necesitaban nuevos endpoints para los "
    "clientes.",
    italic=True, size=10)

figura(doc, "Diagrama de Gantt con la planificacion temporal del proyecto")

doc.add_page_break()


# =================================================================
#  2. DISENO E IMPLEMENTACION DEL PROYECTO
# =================================================================
heading1(doc, "2. Diseno e implementacion del proyecto")


# ---- 2.1 Plan de empresa ----
heading2(doc, "2.1. Plan de empresa")

body(doc,
    "ReciApp se enmarca en una empresa ficticia denominada EcoTech Solutions "
    "S.L., una startup tecnologica especializada en soluciones digitales para "
    "la sostenibilidad medioambiental.")

heading3(doc, "Datos de la empresa")
tabla(doc,
    ["Campo", "Valor"],
    [
        ["Nombre", "EcoTech Solutions S.L."],
        ["Sector", "Tecnologia medioambiental (GreenTech)"],
        ["Actividad", "Desarrollo de aplicaciones para la gestion del reciclaje"],
        ["Sede", "Madrid, Espana"],
        ["Plantilla", "5 empleados (1 CEO, 2 desarrolladores, 1 disenador UX, 1 sysadmin)"],
        ["Clientes objetivo", "Ayuntamientos, comunidades de vecinos, centros educativos, empresas con politicas ESG"],
    ],
    col_widths=[3, 11.5]
)

heading3(doc, "Modelo de negocio")
body(doc,
    "La empresa ofrece ReciApp como servicio (SaaS) a municipios y "
    "organizaciones que deseen implantar un sistema de reciclaje con "
    "trazabilidad. Los ingresos provienen de licencias mensuales por "
    "numero de usuarios activos y de la instalacion de terminales de "
    "reciclaje en puntos estrategicos.")

heading3(doc, "Propuesta de valor")
bullet(doc, "Seguimiento individual del impacto ambiental de cada ciudadano/empleado.")
bullet(doc, "Gamificacion: rankings, acumulacion de CO2 reducido, TAP como identificacion personal.")
bullet(doc, "Multiplataforma: accesible desde movil, escritorio y terminales fisicos.")
bullet(doc, "Datos abiertos para que las organizaciones midan su huella ecologica.")


# ---- 2.2 Arquitectura de la aplicacion ----
heading2(doc, "2.2. Arquitectura de la aplicacion")

body(doc,
    "ReciApp sigue una arquitectura cliente-servidor de tres capas: "
    "presentacion (los tres clientes), logica de negocio (API REST con "
    "Spring Boot) y persistencia (MySQL). Los tres clientes se comunican "
    "exclusivamente con la API REST mediante peticiones HTTP en el puerto "
    "3000, utilizando JSON como formato de intercambio de datos.")

code_block(doc,
"""  +-------------------+     HTTP :3000      +--------------------+    JDBC     +----------+
  |   App Movil       | ------------------> |                    | ---------> |          |
  |  (React Native)   |                     |   API REST         |            |  MySQL   |
  +-------------------+     HTTP :3000      |  (Spring Boot)     |            |  8.x     |
  |  App Escritorio   | ------------------> |                    |            |          |
  |    (JavaFX)       |                     |  Puerto: 3000      |            | Puerto:  |
  +-------------------+     HTTP :3000      |  Host: 0.0.0.0     |            |  3306    |
  |   Terminal        | ------------------> |                    |            |          |
  |   (Python)        |                     +--------------------+            +----------+
  +-------------------+""")

body(doc,
    "La Figura 2 muestra el diagrama de arquitectura del sistema. "
    "Todos los componentes se comunican a traves de HTTP en texto plano. "
    "Se evaluo la implementacion de HTTPS con certificados autofirmados, "
    "pero se descarto por incompatibilidad con Expo Go, tal y como se "
    "detalla en la seccion 1.4.")

figura(doc, "Diagrama de arquitectura del sistema (API central + 3 clientes + MySQL)")

heading3(doc, "Flujo de autenticacion (JWT)")
body(doc,
    "El sistema utiliza JSON Web Tokens (JWT) para la autenticacion "
    "stateless. El flujo es el siguiente:")

bullet(doc, "El cliente envia credenciales (nombre + contrasena) mediante POST a /api/usuarios/login.")
bullet(doc, "El servidor valida la contrasena con BCrypt (strength 10) y genera un token JWT firmado con HMAC-SHA256.")
bullet(doc, "El token contiene los claims: userId, nombre y permisos, con una validez de 7 dias.")
bullet(doc, "En cada peticion posterior, el cliente envia el token en la cabecera Authorization: Bearer <token>.")
bullet(doc, "El filtro JwtAuthFilter intercepta cada peticion, valida el token y establece el SecurityContext de Spring.")

figura(doc, "Diagrama de secuencia del flujo de autenticacion JWT")


# ---- 2.3 Roles de usuario ----
heading2(doc, "2.3. Roles de usuario")

body(doc,
    "El sistema define dos roles de usuario almacenados en la columna "
    "Permisos de la tabla Usuarios, mas un modo invitado exclusivo de "
    "la aplicacion movil.")

heading3(doc, "Rol: Cliente")
body(doc,
    "Es el rol por defecto asignado a todo usuario nuevo al registrarse. "
    "Un usuario con rol cliente puede:")

bullet(doc, "Registrarse e iniciar sesion en cualquiera de las plataformas.")
bullet(doc, "Consultar el catalogo completo de productos reciclables.")
bullet(doc, "Visualizar su historial personal de reciclaje y emisiones acumuladas.")
bullet(doc, "Solicitar y renovar su TAP (Token de Autenticacion Personal) de 6 digitos.")
bullet(doc, "Modificar su nombre de usuario y contrasena.")
bullet(doc, "Eliminar su cuenta (requiere confirmar con contrasena).")
bullet(doc, "Reciclar productos en terminales fisicos usando su TAP.")

heading3(doc, "Rol: Administrador")
body(doc,
    "Tiene todos los permisos del cliente, mas las siguientes "
    "funcionalidades de administracion:")

bullet(doc, "Gestion completa (CRUD) de usuarios a traves de la API (/api/admin/usuarios) y de la app de escritorio.")
bullet(doc, "Gestion completa (CRUD) de productos: crear, listar, modificar y eliminar, incluyendo imagenes en Base64.")
bullet(doc, "Gestion completa (CRUD) de transacciones de reciclaje.")
bullet(doc, "Visualizacion de graficos de emisiones de todos los usuarios (app de escritorio).")
bullet(doc, "Operacion del terminal de reciclaje en puntos fisicos.")

heading3(doc, "Modo Invitado (solo app movil)")
body(doc,
    "La aplicacion movil permite un modo invitado para uso sin conexion o "
    "sin registro. El usuario invitado puede explorar el catalogo de "
    "productos con datos de ejemplo almacenados localmente, pero no tiene "
    "acceso a funcionalidades que requieran autenticacion (historial, TAP, "
    "gestion de perfil). Se muestra un banner informativo en las pantallas "
    "con funcionalidad restringida.")

heading3(doc, "Matriz de permisos")
body(doc,
    "La Tabla 1 resume los permisos de cada rol en el sistema. "
    "La autorizacion se implementa tanto a nivel de API (anotaciones "
    "@PreAuthorize en Spring Security) como a nivel de interfaz "
    "(visibilidad condicional de elementos).")

tabla(doc,
    ["Accion", "Cliente", "Admin", "Invitado"],
    [
        ["Registro / Login", "Si", "Si", "No"],
        ["Ver catalogo de productos", "Si", "Si", "Si (ejemplo)"],
        ["Ver historial propio", "Si", "Si", "No"],
        ["Ver / solicitar TAP", "Si", "Si", "No"],
        ["Modificar nombre / contrasena", "Si", "Si", "No"],
        ["Eliminar cuenta propia", "Si", "Si", "No"],
        ["CRUD Usuarios (admin)", "No", "Si", "No"],
        ["CRUD Productos (admin)", "No", "Si", "No"],
        ["CRUD Transacciones (admin)", "No", "Si", "No"],
        ["Operar terminal de reciclaje", "No", "Si", "No"],
        ["Graficos de emisiones", "No", "Si", "No"],
    ],
    col_widths=[6, 2.5, 2.5, 2.5]
)


# ---- 2.4 Mapa de navegacion ----
heading2(doc, "2.4. Mapa de navegacion")

heading3(doc, "App movil (React Native / Expo)")
body(doc,
    "La navegacion de la aplicacion movil esta basada en Expo Router "
    "(file-based routing). La Figura 4 muestra el mapa de navegacion "
    "completo:")

code_block(doc,
"""                        +--------------------+
                        |    _layout.tsx     |
                        |  (Auth Guard +     |
                        |   Conexion Check)  |
                        +--------+-----------+
                                 |
              +------------------+------------------+
              |                                     |
    +---------v----------+              +-----------v-----------+
    |     (auth)/        |              |       (tabs)/         |
    |  login.tsx         |              | +---+----+-------+    |
    |  register.tsx      |              | | P | H  |  Per  |    |
    +--------------------+              | | r | is |  fil  |    |
                                        | | o | t  |       |    |
    +--------------------+              | | d | or |       |    |
    | connection-modal   |              | | . | .  |       |    |
    |  (Modal offline)   |              | +---+----+-------+    |
    +--------------------+              +-----------------------+

    Leyenda: Prod = Productos, Hist = Historial, Perfil""")

figura(doc, "Mapa de navegacion de la app movil")

heading3(doc, "App de escritorio (JavaFX)")
body(doc,
    "La aplicacion de escritorio sigue un flujo lineal de ventanas "
    "FXML gestionadas por la clase StartWin:")

code_block(doc,
"""  +----------+     +-----------+     +-------------------------+
  |  Login   | --> |  Registro | --> |   Dashboard Principal   |
  +----------+     +-----------+     |  +---------+---------+  |
       |                             |  | Info    | Admin   |  |
       +---->------------------------+  | Person. | (CRUD)  |  |
                                     |  +---------+---------+  |
                                     |  | Grafico emisiones |  |
                                     |  +-------------------+  |
                                     +--------+----------------+
                                              |
                    +------------+------------+------------+
                    |            |            |            |
               Ajustes   Cambiar Pass   NuevoProducto  NuevoUser ...""")

figura(doc, "Mapa de navegacion de la app de escritorio")

heading3(doc, "Terminal Python")
body(doc,
    "El terminal sigue un flujo secuencial en bucle:")

code_block(doc,
"""  Inicio --> Verificar API --> Login operador --> [BUCLE]:
      Escanear codigo de barras --> Buscar producto -->
      Introducir TAP --> Verificar usuario -->
      Confirmar --> Registrar reciclaje --> Mostrar resultado""")


# ---- 2.5 Base de datos ----
heading2(doc, "2.5. Base de datos")

heading3(doc, "2.5.1. Diagrama Entidad-Relacion")

body(doc,
    "La base de datos reciInventario_db utiliza MySQL 8.x con el motor "
    "de almacenamiento InnoDB y codificacion utf8mb4 (soporte completo "
    "de Unicode). El modelo consta de tres entidades principales con las "
    "siguientes relaciones:")

code_block(doc,
"""      +------------------+             +------------------+
      |     USUARIOS     |             |    PRODUCTOS     |
      +------------------+             +------------------+
      | PK Id_Usuario    |             | PK Tipo          |
      |    Nombre (UQ)   |             | PK Numero_barras |
      |    Hash_Contras. |             |    Nombre        |
      |    Permisos      |             |    Emisiones_Red.|
      |    Emisiones_Red.|             |    Material      |
      |    TAP           |             |    Imagen        |
      +--------+---------+             +--------+---------+
               | 1                              | 1
               |                                |
               | N                              | N
      +--------+--------------------------------+---------+
      |                   RECICLA                         |
      +---------------------------------------------------+
      | PK,FK Id_Usuario                                  |
      | PK,FK Tipo                                        |
      | PK,FK Numero_barras                               |
      | PK    Fecha                                       |
      | PK    Hora                                        |
      +---------------------------------------------------+""")

body(doc, "Relaciones del modelo:")
bullet(doc, "Usuarios - Recicla: relacion 1:N. Un usuario puede tener muchos registros de reciclaje.")
bullet(doc, "Productos - Recicla: relacion 1:N. Un producto puede aparecer en muchos registros de reciclaje.")
bullet(doc, "Recicla: tabla intermedia con clave primaria compuesta de cinco columnas (Id_Usuario, Tipo, Numero_barras, Fecha, Hora), lo que permite registrar multiples reciclajes del mismo producto por el mismo usuario en instantes diferentes.")
bullet(doc, "Ambas claves foraneas tienen ON DELETE CASCADE: al eliminar un usuario o producto, se eliminan automaticamente sus registros de reciclaje.")

figura(doc, "Modelo Entidad-Relacion generado desde MySQL Workbench o dbdiagram.io")

heading3(doc, "2.5.2. Paso a tablas")

body(doc, "Tabla: Usuarios", bold=True)
tabla(doc,
    ["Columna", "Tipo de dato", "Restricciones", "Descripcion"],
    [
        ["Id_Usuario", "INT", "PK, AUTO_INCREMENT", "Identificador unico del usuario"],
        ["Nombre", "VARCHAR(50)", "NOT NULL, UNIQUE", "Nombre de usuario para inicio de sesion"],
        ["Hash_Contrasena", "VARCHAR(100)", "NOT NULL", "Hash BCrypt de la contrasena (strength 10)"],
        ["Permisos", "VARCHAR(15)", "DEFAULT 'cliente'", "Rol del usuario: 'cliente' o 'administrador'"],
        ["Emisiones_Reducidas", "FLOAT", "DEFAULT 0", "Kilogramos de CO2 reducidos acumulados"],
        ["TAP", "INT", "DEFAULT NULL", "Token de Autenticacion Personal (6 digitos, unico)"],
    ],
    col_widths=[3.5, 2.8, 3.5, 4.5]
)

body(doc, "Tabla: Productos", bold=True)
body(doc,
    "Utiliza una clave primaria compuesta por el tipo de codigo de barras "
    "y el numero de barras, lo que permite gestionar distintos formatos de "
    "codigos (EAN13, UPC, etc.).")
tabla(doc,
    ["Columna", "Tipo de dato", "Restricciones", "Descripcion"],
    [
        ["Tipo", "VARCHAR(10)", "PK (compuesta)", "Formato del codigo de barras (p.ej. EAN13)"],
        ["Numero_barras", "BIGINT", "PK (compuesta)", "Numero del codigo de barras"],
        ["Nombre", "VARCHAR(50)", "--", "Nombre comercial del producto"],
        ["Emisiones_Reducibles", "FLOAT", "--", "kg de CO2 que se reducen al reciclar el producto"],
        ["Material", "VARCHAR(15)", "--", "Material principal (PET, Vidrio, Aluminio, Brick...)"],
        ["Imagen", "LONGTEXT", "--", "Imagen del producto codificada en Base64"],
    ],
    col_widths=[3.5, 2.8, 3.5, 4.5]
)

body(doc, "Tabla: Recicla (historial de reciclaje)", bold=True)
body(doc,
    "Registra cada accion de reciclaje. Su clave primaria compuesta de cinco "
    "columnas garantiza la unicidad temporal de cada registro.")
tabla(doc,
    ["Columna", "Tipo de dato", "Restricciones", "Descripcion"],
    [
        ["Id_Usuario", "INT", "PK, FK -> Usuarios", "Usuario que realizo el reciclaje"],
        ["Tipo", "VARCHAR(10)", "PK, FK -> Productos", "Tipo de codigo del producto"],
        ["Numero_barras", "BIGINT", "PK, FK -> Productos", "Codigo de barras del producto"],
        ["Fecha", "DATE", "PK", "Fecha en que se realizo el reciclaje"],
        ["Hora", "TIME", "PK", "Hora exacta del reciclaje"],
    ],
    col_widths=[3.5, 2.8, 3.5, 4.5]
)

heading3(doc, "2.5.3. Datos de ejemplo")
body(doc,
    "El script de inicializacion (000_create_database.sql) incluye cinco "
    "productos de ejemplo para facilitar las pruebas del sistema:")
tabla(doc,
    ["Producto", "Codigo EAN13", "Material", "CO2 reducible (kg)"],
    [
        ["Coca-Cola 330ml", "8410076472885", "Aluminio", "0.033"],
        ["Agua Bezoya 1.5L", "8411700000017", "PET", "0.025"],
        ["Leche Hacendado 1L", "8480000291455", "Brick", "0.040"],
        ["Mahou Clasica 330ml", "8410128800505", "Vidrio", "0.045"],
        ["Fanta Naranja 330ml", "5449000000996", "Aluminio", "0.033"],
    ],
    col_widths=[4.5, 3.5, 3, 3]
)

heading3(doc, "2.5.4. Script SQL de creacion")
code_block(doc,
"""CREATE DATABASE IF NOT EXISTS `reciInventario_db`
  DEFAULT CHARACTER SET utf8mb4
  DEFAULT COLLATE utf8mb4_unicode_ci;

USE `reciInventario_db`;

CREATE TABLE IF NOT EXISTS `Usuarios` (
  `Id_Usuario` INT NOT NULL AUTO_INCREMENT,
  `Nombre` VARCHAR(50) NOT NULL,
  `Hash_Contrasena` VARCHAR(100) NOT NULL,
  `Permisos` VARCHAR(15) DEFAULT 'cliente',
  `Emisiones_Reducidas` FLOAT DEFAULT 0,
  `TAP` INT DEFAULT NULL,
  PRIMARY KEY (`Id_Usuario`),
  UNIQUE KEY `uk_nombre` (`Nombre`)
) ENGINE=InnoDB DEFAULT CHARSET=utf8mb4;

CREATE TABLE IF NOT EXISTS `Productos` (
  `Tipo` VARCHAR(10) NOT NULL,
  `Numero_barras` BIGINT NOT NULL,
  `Nombre` VARCHAR(50) DEFAULT NULL,
  `Emisiones_Reducibles` FLOAT DEFAULT NULL,
  `Material` VARCHAR(15) DEFAULT NULL,
  `Imagen` LONGTEXT DEFAULT NULL,
  PRIMARY KEY (`Tipo`, `Numero_barras`)
) ENGINE=InnoDB DEFAULT CHARSET=utf8mb4;

CREATE TABLE IF NOT EXISTS `Recicla` (
  `Id_Usuario` INT NOT NULL,
  `Tipo` VARCHAR(10) NOT NULL,
  `Numero_barras` BIGINT NOT NULL,
  `Fecha` DATE NOT NULL,
  `Hora` TIME NOT NULL,
  PRIMARY KEY (`Id_Usuario`,`Tipo`,`Numero_barras`,`Fecha`,`Hora`),
  CONSTRAINT `fk_recicla_usuario`
    FOREIGN KEY (`Id_Usuario`) REFERENCES `Usuarios` (`Id_Usuario`)
    ON DELETE CASCADE,
  CONSTRAINT `fk_recicla_producto`
    FOREIGN KEY (`Tipo`,`Numero_barras`) REFERENCES `Productos` (`Tipo`,`Numero_barras`)
    ON DELETE CASCADE
) ENGINE=InnoDB DEFAULT CHARSET=utf8mb4;""")


# ---- 2.6 Implementacion por modulos ----
heading2(doc, "2.6. Implementacion por modulos")

body(doc,
    "A continuacion se detalla como el proyecto cubre cada uno de los "
    "modulos del ciclo formativo de DAM.")

# --- Acceso a Datos ---
heading3(doc, "2.6.1. Acceso a Datos")

body(doc,
    "El modulo de Acceso a Datos se cubre principalmente a traves de la "
    "API REST desarrollada con Spring Boot y Spring Data JPA:")

bullet(doc, "Utilizacion de un ORM (Hibernate/JPA): las entidades Java (Usuario, Producto, Recicla) se mapean directamente a las tablas de MySQL mediante anotaciones JPA (@Entity, @Table, @Id, @Column, @ManyToOne, @EmbeddedId).")
bullet(doc, "Repositorios Spring Data JPA: las interfaces UsuarioRepository, ProductoRepository y ReciclaRepository extienden JpaRepository, proporcionando operaciones CRUD automaticas y consultas derivadas del nombre del metodo (findByNombre, findByTap, existsByNombre).")
bullet(doc, "API REST como metodo de acceso preferente: todos los clientes acceden a la base de datos exclusivamente a traves de la API REST (endpoints HTTP), cumpliendo con la recomendacion del proyecto de utilizar una API REST como metodo mas versatil y efectivo.")
bullet(doc, "Base de datos relacional remota: MySQL se ejecuta en el servidor EC2, siendo accedida por la API mediante JDBC (driver mysql-connector-j).")
bullet(doc, "DTOs para transferencia de datos: se utilizan objetos DTO (LoginRequest, AuthResponse, ProductoDto, HistorialDto, etc.) para desacoplar la capa de presentacion de las entidades de persistencia.")

body(doc, "Estructura de la capa de acceso a datos:", bold=True)

code_block(doc,
"""api-spring/src/main/java/com/reciapp/api/
+-- entity/
|   +-- Usuario.java          (JPA @Entity -> tabla Usuarios)
|   +-- Producto.java         (JPA @Entity -> tabla Productos)
|   +-- ProductoId.java       (@EmbeddedId, clave compuesta)
|   +-- Recicla.java          (JPA @Entity -> tabla Recicla)
|   +-- ReciclaId.java        (@EmbeddedId, clave compuesta)
+-- repository/
|   +-- UsuarioRepository.java    (extends JpaRepository)
|   +-- ProductoRepository.java   (extends JpaRepository)
|   +-- ReciclaRepository.java    (extends JpaRepository)
+-- service/
|   +-- UsuarioService.java       (logica de negocio usuarios)
|   +-- ProductoService.java      (logica de negocio productos)
|   +-- ReciclaService.java       (logica de negocio reciclaje)
|   +-- AdminService.java         (operaciones de administracion)""")

figura(doc, "Captura de una entidad JPA en el IDE mostrando las anotaciones de mapeo")

# --- Desarrollo de Interfaces ---
heading3(doc, "2.6.2. Desarrollo de Interfaces")

body(doc,
    "El modulo de Desarrollo de Interfaces se cubre con la aplicacion de "
    "escritorio desarrollada en JavaFX 25:")

bullet(doc, "Interfaz grafica multiplataforma con JavaFX y FXML: la aplicacion cuenta con 12 vistas FXML que definen las interfaces de forma declarativa, separando la estructura visual (FXML) de la logica (controladores Java).")
bullet(doc, "Estilizacion mediante CSS: se utiliza un archivo styles.css para personalizar la apariencia de los componentes JavaFX.")
bullet(doc, "Formularios para gestion de datos: se han implementado formularios completos para crear y modificar usuarios, productos y transacciones, con validacion de campos.")
bullet(doc, "Informes y graficos: la pestana de graficos utiliza la libreria ChartFX para generar graficos de barras con las emisiones reducidas por cada usuario del sistema.")
bullet(doc, "Navegacion entre ventanas: la clase StartWin gestiona la transicion entre pantallas mediante carga dinamica de archivos FXML.")

body(doc, "Vistas FXML implementadas:", bold=True)
tabla(doc,
    ["Vista FXML", "Controlador", "Funcion"],
    [
        ["loginStart_win.fxml", "LoginController", "Inicio de sesion"],
        ["singUp_win.fxml", "SingUpController", "Registro de nuevo usuario"],
        ["main_win.fxml", "MainController", "Dashboard con 3 pestanas (info, admin, grafico)"],
        ["settings_win.fxml", "SettingsController", "Configuracion del usuario"],
        ["changePasswd.fxml", "ChangePasswd", "Cambio de contrasena"],
        ["escanear_win.fxml", "Escanear", "Escaneo de codigos de barras"],
        ["newProducto.fxml", "NewProducto", "Crear nuevo producto"],
        ["modProducto.fxml", "ModProducto", "Modificar producto existente"],
        ["newUser.fxml", "NewUser", "Crear nuevo usuario"],
        ["modUser.fxml", "ModUser", "Modificar usuario existente"],
        ["newTransaccion.fxml", "NewTransaccion", "Crear nueva transaccion"],
        ["modTransaccion.fxml", "ModTransaccion", "Modificar transaccion"],
    ],
    col_widths=[4, 3.5, 6]
)

figura(doc, "Interfaz principal de la app de escritorio (dashboard con pestanas)")
figura(doc, "Formulario de creacion/modificacion de un producto (vista FXML)")

# --- Programacion de Servicios y Procesos ---
heading3(doc, "2.6.3. Programacion de Servicios y Procesos")

body(doc,
    "El modulo de Programacion de Servicios y Procesos se cubre con los "
    "siguientes aspectos del proyecto:")

bullet(doc, "Comunicacion en red mediante HTTP: todos los clientes se comunican con la API REST a traves de peticiones HTTP (GET, POST, PUT, DELETE) utilizando sockets TCP de forma transparente.")
bullet(doc, "Servidor web concurrente: Spring Boot gestiona multiples hilos de ejecucion para atender peticiones simultaneas de los distintos clientes (movil, escritorio y terminal).")
bullet(doc, "HttpClient multihilo en JavaFX: la aplicacion de escritorio utiliza java.net.http.HttpClient, que internamente usa un pool de hilos para gestionar conexiones HTTP asincronas.")
bullet(doc, "Peticiones asincronas en React Native: la app movil realiza peticiones HTTP asincronas (fetch/async-await) sin bloquear el hilo principal de la interfaz de usuario.")
bullet(doc, "Sesiones HTTP en Python: el terminal utiliza requests.Session() para mantener la conexion y reutilizar cabeceras (incluyendo el token JWT) entre multiples peticiones.")

body(doc, "Endpoints de la API REST:", bold=True)
body(doc, "Endpoints publicos (sin autenticacion):", bold=True, size=10)
tabla(doc,
    ["Metodo", "Endpoint", "Descripcion"],
    [
        ["GET", "/api/health", "Comprobacion de estado del servidor"],
        ["POST", "/api/usuarios/login", "Autenticacion (devuelve JWT)"],
        ["POST", "/api/usuarios/register", "Registro de nuevo usuario"],
    ],
    col_widths=[2, 5, 7.5]
)

body(doc, "Endpoints protegidos (requieren JWT):", bold=True, size=10)
tabla(doc,
    ["Metodo", "Endpoint", "Descripcion"],
    [
        ["GET", "/api/usuarios/profile", "Perfil del usuario autenticado"],
        ["PUT", "/api/usuarios/{id}/nombre", "Actualizar nombre"],
        ["PUT", "/api/usuarios/{id}/password", "Cambiar contrasena"],
        ["PUT", "/api/usuarios/{id}/tap", "Solicitar/renovar TAP"],
        ["DELETE", "/api/usuarios/{id}", "Eliminar cuenta"],
        ["GET", "/api/usuarios/by-tap/{tap}", "Buscar usuario por TAP"],
        ["GET", "/api/productos", "Listar todos los productos"],
        ["GET", "/api/productos/barcode/{code}", "Buscar producto por codigo de barras"],
        ["GET", "/api/historial", "Historial de reciclaje del usuario"],
        ["POST", "/api/historial", "Registrar nuevo reciclaje"],
    ],
    col_widths=[2, 5, 7.5]
)

body(doc, "Endpoints de administracion (requieren JWT + rol administrador):", bold=True, size=10)
tabla(doc,
    ["Metodo", "Endpoint", "Descripcion"],
    [
        ["GET", "/api/admin/usuarios", "Listar todos los usuarios"],
        ["POST", "/api/admin/usuarios", "Crear usuario"],
        ["PUT", "/api/admin/usuarios/{id}", "Modificar usuario"],
        ["DELETE", "/api/admin/usuarios/{id}", "Eliminar usuario"],
        ["GET", "/api/admin/productos", "Listar todos los productos"],
        ["POST", "/api/admin/productos", "Crear producto"],
        ["PUT", "/api/admin/productos/{tipo}/{barras}", "Modificar producto"],
        ["DELETE", "/api/admin/productos/{tipo}/{barras}", "Eliminar producto"],
        ["GET", "/api/admin/transacciones", "Listar transacciones"],
        ["POST", "/api/admin/transacciones", "Crear transaccion"],
        ["DELETE", "/api/admin/transacciones", "Eliminar transaccion"],
    ],
    col_widths=[2, 5, 7.5]
)

# --- Multimedia y Dispositivos Moviles ---
heading3(doc, "2.6.4. Programacion Multimedia y Dispositivos Moviles")

body(doc,
    "El modulo de Programacion Multimedia y Dispositivos Moviles se cubre "
    "con la aplicacion movil desarrollada en React Native 0.81.5 con Expo "
    "SDK 54 y TypeScript:")

bullet(doc, "Interactividad del usuario: la app ofrece navegacion por pestanas (Productos, Historial, Perfil), pull-to-refresh, busqueda interactiva, filtros por material, modales, formularios y animaciones de transicion.")
bullet(doc, "Intercambio de informacion: se comunica con la API REST para obtener y enviar datos en tiempo real, con soporte para modo offline con datos cacheados en AsyncStorage.")
bullet(doc, "Usabilidad: disenada siguiendo principios de Material Design con NativeWind (Tailwind CSS) para un estilo coherente, Ionicons para iconografia y React Native Paper para componentes estandar.")

body(doc, "Estructura de la app movil:", bold=True)
code_block(doc,
"""app/
+-- _layout.tsx              <-- Layout raiz: auth guard + verificacion de conexion
+-- index.tsx                <-- Punto de entrada
+-- connection-modal.tsx     <-- Modal de modo offline
+-- (auth)/
|   +-- _layout.tsx
|   +-- login.tsx            <-- Pantalla de inicio de sesion
|   +-- register.tsx         <-- Pantalla de registro
+-- (tabs)/
    +-- _layout.tsx          <-- Tab navigator (3 pestanas)
    +-- productos.tsx        <-- Catalogo de productos
    +-- productos/[barcode].tsx  <-- Detalle de producto
    +-- historial.tsx        <-- Historial de reciclaje
    +-- perfil.tsx           <-- Perfil del usuario

store/
+-- authStore.ts             <-- Token JWT, usuario, login/logout, modo invitado
+-- connectionStore.ts       <-- Estado de conexion con la API
+-- productStore.ts          <-- Cache de productos
+-- recycleStore.ts          <-- Historial de reciclaje

services/
+-- api.ts                   <-- Interfaz ApiService (contrato)
+-- api.client.ts            <-- RealApiService (HTTP)
+-- api.offline.ts           <-- OfflineApiService (datos simulados)
+-- index.ts                 <-- Factoria de servicios""")

figura(doc, "Pantalla de login de la app movil")
figura(doc, "Pantalla de productos con filtro de materiales activo")
figura(doc, "Pantalla de historial mostrando reciclajes y total CO2 reducido")
figura(doc, "Pantalla de perfil con estadisticas y TAP")
figura(doc, "Modal de conexion offline")

heading3(doc, "2.6.5. Terminal de Reciclaje (Python)")

body(doc,
    "Adicionalmente, se ha desarrollado un terminal de reciclaje en Python "
    "para puntos fisicos. Este componente complementa el sistema permitiendo "
    "el registro de reciclajes mediante escaneo de codigos de barras e "
    "identificacion por TAP.")

bullet(doc, "Utiliza la libreria requests para comunicacion HTTP con la API REST.")
bullet(doc, "Interfaz de terminal con colores ANSI para mejorar la experiencia visual.")
bullet(doc, "Configuracion mediante variables de entorno (RECIAPP_API_URL, RECIAPP_ADMIN_USER, RECIAPP_ADMIN_PASS).")
bullet(doc, "Flujo de 3 pasos: escanear producto -> introducir TAP -> confirmar reciclaje.")

figura(doc, "Terminal de reciclaje: banner de bienvenida y flujo completo")

doc.add_page_break()


# =================================================================
#  3. FASE DE PRUEBAS
# =================================================================
heading1(doc, "3. Fase de pruebas")

body(doc,
    "Se ha disenado una bateria de pruebas funcionales para verificar el "
    "correcto funcionamiento de todos los componentes del sistema. Las "
    "pruebas se han realizado de forma manual, comprobando cada flujo de "
    "la aplicacion desde los distintos clientes.")

heading2(doc, "3.1. Pruebas de la API REST")

tabla(doc,
    ["ID", "Prueba", "Entrada", "Resultado esperado", "Estado"],
    [
        ["P-01", "Health check", "GET /api/health", '{"status":"ok"}', "OK"],
        ["P-02", "Registro de usuario", "POST /api/usuarios/register con nombre y contrasena", "201 Created + token JWT", "OK"],
        ["P-03", "Registro duplicado", "POST /api/usuarios/register con nombre existente", "409 Conflict", "OK"],
        ["P-04", "Login correcto", "POST /api/usuarios/login con credenciales validas", "200 OK + token JWT", "OK"],
        ["P-05", "Login incorrecto", "POST /api/usuarios/login con contrasena erronea", "401 Unauthorized", "OK"],
        ["P-06", "Acceso sin token", "GET /api/productos sin cabecera Authorization", "401 Unauthorized", "OK"],
        ["P-07", "Acceso con token", "GET /api/productos con Bearer token valido", "200 OK + lista de productos", "OK"],
        ["P-08", "Solicitar TAP", "PUT /api/usuarios/{id}/tap", "200 OK + TAP de 6 digitos", "OK"],
        ["P-09", "Buscar por TAP", "GET /api/usuarios/by-tap/{tap}", "200 OK + datos del usuario", "OK"],
        ["P-10", "Registrar reciclaje", "POST /api/historial con idUsuario, tipo y numeroBarras", "201 Created + emisiones acumuladas", "OK"],
        ["P-11", "Admin sin permisos", "GET /api/admin/usuarios con token de cliente", "403 Forbidden", "OK"],
        ["P-12", "Admin con permisos", "GET /api/admin/usuarios con token de admin", "200 OK + lista de usuarios", "OK"],
    ],
    col_widths=[1.2, 3, 4, 4, 1.3]
)

heading2(doc, "3.2. Pruebas de la aplicacion movil")

tabla(doc,
    ["ID", "Prueba", "Procedimiento", "Resultado esperado", "Estado"],
    [
        ["M-01", "Login exitoso", "Introducir credenciales validas y pulsar Iniciar Sesion", "Redireccion a pantalla de Productos", "OK"],
        ["M-02", "Login fallido", "Introducir contrasena incorrecta", "Alert con mensaje de error", "OK"],
        ["M-03", "Registro", "Rellenar formulario de registro con datos nuevos", "Cuenta creada y redireccion a Productos", "OK"],
        ["M-04", "Ver productos", "Navegar a pestana Productos", "Lista de productos cargada desde la API", "OK"],
        ["M-05", "Filtrar por material", "Pulsar chip de material (Aluminio)", "Solo se muestran productos de Aluminio", "OK"],
        ["M-06", "Buscar producto", "Escribir nombre en la barra de busqueda", "Lista filtrada en tiempo real", "OK"],
        ["M-07", "Ver historial", "Navegar a pestana Historial", "Lista de reciclajes con total de CO2", "OK"],
        ["M-08", "Ver/solicitar TAP", "Ir a Perfil -> Ver TAP -> Solicitar TAP", "Modal muestra TAP de 6 digitos", "OK"],
        ["M-09", "Cambiar nombre", "Perfil -> nuevo nombre + contrasena actual", "Nombre actualizado correctamente", "OK"],
        ["M-10", "Modo offline", "Abrir app sin conexion -> Modo offline", "App funciona con datos de ejemplo", "OK"],
        ["M-11", "Pull to refresh", "Deslizar hacia abajo en Productos", "Lista se recarga desde la API", "OK"],
    ],
    col_widths=[1.2, 3, 4.5, 3.5, 1.3]
)

heading2(doc, "3.3. Pruebas de la aplicacion de escritorio")

tabla(doc,
    ["ID", "Prueba", "Procedimiento", "Resultado esperado", "Estado"],
    [
        ["E-01", "Login", "Introducir credenciales en la ventana de login", "Acceso al dashboard principal", "OK"],
        ["E-02", "Tab Admin visible", "Login con cuenta de administrador", "Pestana de administracion visible", "OK"],
        ["E-03", "Tab Admin oculta", "Login con cuenta de cliente", "Pestana de administracion no visible", "OK"],
        ["E-04", "CRUD Usuarios", "Crear, modificar y eliminar un usuario de prueba", "Operaciones reflejadas en la tabla", "OK"],
        ["E-05", "CRUD Productos", "Crear, modificar y eliminar un producto de prueba", "Operaciones reflejadas en la tabla", "OK"],
        ["E-06", "Grafico emisiones", "Navegar a la pestana de graficos", "Grafico de barras con emisiones por usuario", "OK"],
    ],
    col_widths=[1.2, 3, 4.5, 3.5, 1.3]
)

heading2(doc, "3.4. Pruebas del terminal Python")

tabla(doc,
    ["ID", "Prueba", "Procedimiento", "Resultado esperado", "Estado"],
    [
        ["T-01", "Conexion API", "Ejecutar terminal con API activa", "Mensaje: Conexion con la API establecida", "OK"],
        ["T-02", "Login operador", "Autenticar con credenciales de admin", "Mensaje: Terminal autenticado correctamente", "OK"],
        ["T-03", "Escanear producto", "Introducir codigo 8410076472885", "Producto Coca-Cola 330ml encontrado", "OK"],
        ["T-04", "Producto inexistente", "Introducir codigo inventado", "Error: Producto no encontrado", "OK"],
        ["T-05", "TAP valido", "Introducir TAP de un usuario existente", "Usuario verificado con nombre y emisiones", "OK"],
        ["T-06", "TAP invalido", "Introducir TAP inexistente", "Error: No se encontro usuario con ese TAP", "OK"],
        ["T-07", "Reciclaje completo", "Escanear + TAP + confirmar", "Reciclaje registrado, emisiones acumuladas mostradas", "OK"],
    ],
    col_widths=[1.2, 3, 4, 4, 1.3]
)

heading2(doc, "3.5. Pruebas de integracion")

tabla(doc,
    ["ID", "Prueba", "Procedimiento", "Resultado esperado", "Estado"],
    [
        ["I-01", "Terminal -> Movil", "Registrar reciclaje en terminal y verificar en app movil", "El historial de la app movil refleja el nuevo reciclaje", "OK"],
        ["I-02", "Escritorio -> Movil", "Crear producto en escritorio y verificar en app movil", "El nuevo producto aparece en la lista de productos movil", "OK"],
        ["I-03", "Movil -> Terminal", "Solicitar TAP en app movil y usarlo en terminal", "El terminal reconoce el TAP y permite reciclar", "OK"],
        ["I-04", "Concurrencia", "Usar app movil y escritorio simultaneamente", "Ambas apps funcionan sin interferencias", "OK"],
    ],
    col_widths=[1.2, 3, 4, 4, 1.3]
)

doc.add_page_break()


# =================================================================
#  4. DOCUMENTACION DE LA APLICACION
# =================================================================
heading1(doc, "4. Documentacion de la aplicacion")


# ---- 4.1 Introduccion ----
heading2(doc, "4.1. Introduccion a la aplicacion")

body(doc,
    "ReciApp es un sistema de gestion de reciclaje multiplataforma que "
    "permite registrar productos reciclados y realizar un seguimiento de "
    "las emisiones de CO2 reducidas por cada usuario. El sistema esta "
    "compuesto por cuatro aplicaciones que se comunican con una API REST "
    "centralizada.")

body(doc,
    "El sistema esta pensado para dos tipos de usuarios: clientes (usuarios "
    "estandar que reciclan y consultan su impacto) y administradores (que "
    "gestionan el catalogo de productos, los usuarios y las transacciones). "
    "Adicionalmente, la app movil permite un modo invitado para explorar "
    "la aplicacion sin necesidad de registro.")


# ---- 4.2 Manual de Instalacion ----
heading2(doc, "4.2. Manual de instalacion")

heading3(doc, "4.2.1. Requisitos previos")
tabla(doc,
    ["Componente", "Requisito"],
    [
        ["Base de datos", "MySQL 8.x instalado y en ejecucion"],
        ["API REST", "JDK 21 (Eclipse Temurin recomendado)"],
        ["App escritorio", "JDK 24, JavaFX 25"],
        ["App movil", "Node.js 20 LTS, Expo CLI, dispositivo Android/iOS o emulador"],
        ["Terminal", "Python 3.10+, libreria requests"],
    ],
    col_widths=[3.5, 11]
)

heading3(doc, "4.2.2. Instalacion de la base de datos")
body(doc, "Ejecutar los siguientes comandos en el servidor MySQL:")
code_block(doc,
"""mysql -u root -p < sql/000_create_database.sql
mysql -u root -p < sql/001_imagen_longtext.sql""")

heading3(doc, "4.2.3. Instalacion de la API REST")
code_block(doc,
"""cd api-spring
./gradlew build                    # Compilar el proyecto
java -jar build/libs/api-0.0.1-SNAPSHOT.jar   # Ejecutar en puerto 3000""")
body(doc,
    "La API escuchara en http://0.0.0.0:3000. Se debe verificar que el "
    "puerto 3000 este abierto en el firewall o Security Group de AWS.")

heading3(doc, "4.2.4. Instalacion de la app movil")
code_block(doc,
"""cd reci_app
npm install                        # Instalar dependencias
npx expo start                     # Iniciar servidor de desarrollo""")
body(doc,
    "Se abrira la interfaz de Expo. Se puede escanear el codigo QR con "
    "la app Expo Go en un dispositivo fisico o pulsar 'a' para abrir en "
    "un emulador Android.")

heading3(doc, "4.2.5. Instalacion de la app de escritorio")
code_block(doc,
"""cd ActualAgain
# Abrir el proyecto en IntelliJ IDEA o ejecutar:
gradle run""")
body(doc,
    "Asegurarse de que el archivo configuration.properties contiene la "
    "URL correcta de la API (api.url=http://52.201.91.206:3000).")

heading3(doc, "4.2.6. Instalacion del terminal Python")
code_block(doc,
"""cd terminal
pip install requests               # Instalar dependencia
python reciclaje_terminal.py       # Ejecutar""")
body(doc,
    "Opcionalmente, se pueden configurar variables de entorno: "
    "RECIAPP_API_URL, RECIAPP_ADMIN_USER y RECIAPP_ADMIN_PASS.")

heading3(doc, "4.2.7. Despliegue con Docker (opcional)")
code_block(doc,
"""# docker-compose.yml
version: '3.8'
services:
  mysql:
    image: mysql:8.0
    container_name: reciapp-db
    environment:
      MYSQL_ROOT_PASSWORD: clase1234
      MYSQL_DATABASE: reciInventario_db
    ports:
      - "3306:3306"
    volumes:
      - mysql_data:/var/lib/mysql
      - ./sql/000_create_database.sql:/docker-entrypoint-initdb.d/01.sql
      - ./sql/001_imagen_longtext.sql:/docker-entrypoint-initdb.d/02.sql
    healthcheck:
      test: ["CMD", "mysqladmin", "ping", "-h", "localhost"]
      interval: 10s
      retries: 5

  api:
    build: ./api-spring
    container_name: reciapp-api
    ports:
      - "3000:3000"
    environment:
      SPRING_DATASOURCE_URL: jdbc:mysql://mysql:3306/reciInventario_db
      SPRING_DATASOURCE_USERNAME: root
      SPRING_DATASOURCE_PASSWORD: clase1234
    depends_on:
      mysql:
        condition: service_healthy

volumes:
  mysql_data:""")

code_block(doc,
"""# api-spring/Dockerfile
FROM eclipse-temurin:21-jre-alpine
WORKDIR /app
COPY build/libs/*.jar app.jar
EXPOSE 3000
ENTRYPOINT ["java", "-jar", "app.jar"]""")

body(doc, "Para levantar todo el sistema con Docker:")
code_block(doc, "docker compose up -d")

figura(doc, "Consola EC2 mostrando la API en ejecucion o docker compose up")


# ---- 4.3 Manual de usuario ----
heading2(doc, "4.3. Manual de usuario")

heading3(doc, "4.3.1. Registro e inicio de sesion")
body(doc,
    "Al abrir la aplicacion (movil o escritorio), se muestra la pantalla "
    "de inicio de sesion. Si el usuario no tiene cuenta, puede pulsar en "
    "'Registrate' para acceder al formulario de registro. Se debe introducir "
    "un nombre de usuario unico y una contrasena. Una vez registrado, el "
    "sistema inicia sesion automaticamente.")

figura(doc, "Pantalla de login y registro de la app movil")

heading3(doc, "4.3.2. Consultar productos")
body(doc,
    "La pestana 'Productos' muestra el catalogo completo de productos "
    "reciclables. Cada tarjeta muestra el nombre del producto, su material, "
    "codigo de barras y la cantidad de CO2 que se reduce al reciclarlo. "
    "Se puede buscar por nombre o codigo de barras usando la barra de "
    "busqueda, y filtrar por tipo de material usando los chips superiores.")

figura(doc, "Pantalla de productos con busqueda y filtros")

heading3(doc, "4.3.3. Consultar historial de reciclaje")
body(doc,
    "La pestana 'Historial' muestra todas las acciones de reciclaje "
    "realizadas por el usuario. En la cabecera se visualiza el total de "
    "CO2 reducido acumulado. Cada entrada del historial incluye el nombre "
    "del producto, su material, la fecha del reciclaje y las emisiones "
    "reducidas.")

heading3(doc, "4.3.4. Gestionar perfil")
body(doc,
    "La pestana 'Perfil' permite al usuario ver sus estadisticas "
    "(emisiones en kg y toneladas), consultar y solicitar su TAP, "
    "cambiar su nombre de usuario o contrasena, y eliminar su cuenta.")

heading3(doc, "4.3.5. Uso del TAP en el terminal de reciclaje")
body(doc,
    "Para reciclar en un punto fisico, el usuario debe tener un TAP "
    "activo (se solicita desde el perfil de la app movil). En el terminal "
    "de reciclaje, el operador escanea el codigo de barras del producto "
    "y el usuario introduce su TAP de 6 digitos. El sistema registra el "
    "reciclaje y muestra la confirmacion con las emisiones acumuladas.")

heading3(doc, "4.3.6. Modo offline (app movil)")
body(doc,
    "Si la app movil no puede conectar con la API, se muestra un modal "
    "con dos opciones: 'Recargar' para reintentar la conexion, o 'Modo "
    "offline' para explorar la app con datos de ejemplo. En modo offline, "
    "las funcionalidades de autenticacion y gestion estan deshabilitadas.")


# ---- 4.4 Manual de administracion ----
heading2(doc, "4.4. Manual de administracion")

body(doc,
    "La administracion del sistema se realiza principalmente desde la "
    "aplicacion de escritorio (JavaFX), que ofrece una interfaz grafica "
    "completa para la gestion de usuarios, productos y transacciones.")

heading3(doc, "4.4.1. Acceso al panel de administracion")
body(doc,
    "Al iniciar sesion con una cuenta de administrador, el dashboard "
    "muestra una pestana adicional de 'Administracion' que no esta "
    "visible para los usuarios con rol cliente.")

heading3(doc, "4.4.2. Gestion de usuarios")
body(doc,
    "La seccion de usuarios muestra una tabla con todos los usuarios "
    "del sistema (ID, nombre, permisos, emisiones, TAP). El administrador "
    "puede crear nuevos usuarios, modificar sus datos (incluyendo cambiar "
    "su rol) y eliminar usuarios. La eliminacion es en cascada: se borran "
    "automaticamente los registros de reciclaje asociados.")

heading3(doc, "4.4.3. Gestion de productos")
body(doc,
    "La seccion de productos muestra el catalogo completo. El administrador "
    "puede crear nuevos productos indicando tipo de codigo, numero de barras, "
    "nombre, material, emisiones reducibles e imagen (en Base64). Tambien "
    "puede modificar cualquier campo de un producto existente o eliminarlo.")

heading3(doc, "4.4.4. Gestion de transacciones")
body(doc,
    "La seccion de transacciones muestra el historial de reciclaje de "
    "todos los usuarios. El administrador puede crear transacciones "
    "manualmente (util para corregir datos o registrar reciclajes "
    "retroactivos) y eliminar transacciones incorrectas.")

heading3(doc, "4.4.5. Graficos de emisiones")
body(doc,
    "La pestana de graficos muestra un grafico de barras con las emisiones "
    "de CO2 reducidas por cada usuario del sistema, permitiendo una "
    "visualizacion comparativa del impacto de cada participante.")

figura(doc, "Panel de administracion con tabla de usuarios")
figura(doc, "Panel de administracion con tabla de productos")
figura(doc, "Grafico de emisiones por usuario")

doc.add_page_break()


# =================================================================
#  5. CONCLUSIONES FINALES
# =================================================================
heading1(doc, "5. Conclusiones finales")

heading2(doc, "5.1. Grado de cumplimiento de los objetivos")

body(doc,
    "Se ha logrado desarrollar un sistema multiplataforma completo y "
    "funcional que cumple con todos los objetivos planteados al inicio "
    "del proyecto:")

bullet(doc, "Se ha implementado una API REST robusta con Spring Boot que gestiona toda la logica de negocio, la autenticacion JWT y la persistencia en MySQL.")
bullet(doc, "Se ha desarrollado una aplicacion movil multiplataforma con React Native y Expo que permite a los usuarios consultar productos, ver su historial y gestionar su perfil.")
bullet(doc, "Se ha creado una aplicacion de escritorio con JavaFX que proporciona funcionalidades de administracion completas, incluyendo graficos de emisiones.")
bullet(doc, "Se ha implementado un terminal de reciclaje en Python para puntos fisicos, utilizando codigos de barras y TAP como metodo de identificacion.")
bullet(doc, "El sistema se ha desplegado exitosamente en un servidor AWS EC2, demostrando su viabilidad en un entorno real.")
bullet(doc, "El proyecto cubre los modulos de Acceso a Datos, Desarrollo de Interfaces, Programacion de Servicios y Procesos, y Programacion Multimedia y Dispositivos Moviles.")


heading2(doc, "5.2. Propuesta de modificaciones y ampliaciones futuras")

body(doc,
    "A continuacion se proponen posibles mejoras y ampliaciones del "
    "sistema para futuras versiones:")

bullet(doc, "HTTPS en produccion: configurar un certificado SSL de Let's Encrypt detras de un proxy Nginx para cifrar todas las comunicaciones.")
bullet(doc, "Sistema de rankings y gamificacion: implementar tablas de clasificacion de usuarios por emisiones reducidas, insignias y logros.")
bullet(doc, "Notificaciones push: avisar a los usuarios de nuevos productos anadidos al catalogo o recordatorios de reciclaje.")
bullet(doc, "Escaneo de codigos de barras desde la app movil: integrar la camara del dispositivo para escanear productos directamente, sin necesidad de un terminal fisico.")
bullet(doc, "Dashboard web de administracion: crear un panel de control web con React o Angular que complemente la app de escritorio.")
bullet(doc, "Analisis de datos: implementar estadisticas avanzadas con graficos temporales, comparativas por materiales y predicciones de impacto.")
bullet(doc, "Internacionalizacion (i18n): traducir la aplicacion a otros idiomas para ampliar su alcance.")
bullet(doc, "Autenticacion con redes sociales: permitir login con Google, Apple o Facebook para facilitar el registro.")
bullet(doc, "Tests automatizados: implementar tests unitarios (JUnit, Jest) y de integracion (Cypress, Detox) para mejorar la calidad del codigo.")

doc.add_page_break()


# =================================================================
#  6. BIBLIOGRAFIA
# =================================================================
heading1(doc, "6. Bibliografia")

body(doc, "Documentacion oficial y recursos utilizados durante el desarrollo del proyecto:", bold=True)

heading3(doc, "Backend y base de datos")
bullet(doc, "Spring Boot Reference Documentation. https://docs.spring.io/spring-boot/docs/current/reference/html/")
bullet(doc, "Spring Security Reference. https://docs.spring.io/spring-security/reference/index.html")
bullet(doc, "Spring Data JPA Reference. https://docs.spring.io/spring-data/jpa/docs/current/reference/html/")
bullet(doc, "MySQL 8.0 Reference Manual. https://dev.mysql.com/doc/refman/8.0/en/")
bullet(doc, "JSON Web Tokens (JWT) Introduction. https://jwt.io/introduction")
bullet(doc, "JJWT - Java JWT. https://github.com/jwtk/jjwt")

heading3(doc, "Aplicacion movil")
bullet(doc, "React Native Documentation. https://reactnative.dev/docs/getting-started")
bullet(doc, "Expo Documentation. https://docs.expo.dev/")
bullet(doc, "Expo Router Documentation. https://docs.expo.dev/router/introduction/")
bullet(doc, "NativeWind Documentation. https://www.nativewind.dev/")
bullet(doc, "Zustand - State Management. https://docs.pmnd.rs/zustand/getting-started/introduction")
bullet(doc, "React Native Paper. https://callstack.github.io/react-native-paper/")

heading3(doc, "Aplicacion de escritorio")
bullet(doc, "OpenJFX Documentation. https://openjfx.io/javadoc/21/")
bullet(doc, "JavaFX CSS Reference Guide. https://openjfx.io/javadoc/21/javafx.graphics/javafx/scene/doc-files/cssref.html")
bullet(doc, "Gson User Guide. https://github.com/google/gson/blob/master/UserGuide.md")

heading3(doc, "Terminal y Python")
bullet(doc, "Python Requests Library. https://docs.python-requests.org/en/latest/")
bullet(doc, "Python 3 Documentation. https://docs.python.org/3/")

heading3(doc, "Despliegue e infraestructura")
bullet(doc, "AWS EC2 User Guide. https://docs.aws.amazon.com/ec2/")
bullet(doc, "Docker Documentation. https://docs.docker.com/")
bullet(doc, "Docker Compose Documentation. https://docs.docker.com/compose/")

heading3(doc, "Herramientas de desarrollo")
bullet(doc, "Git Documentation. https://git-scm.com/doc")
bullet(doc, "Postman Learning Center. https://learning.postman.com/docs/getting-started/introduction/")
bullet(doc, "IntelliJ IDEA Documentation. https://www.jetbrains.com/idea/documentation/")

body(doc, "")
body(doc, "")


# =====================================================================
# GUARDAR DOCUMENTO
# =====================================================================

output_path = os.path.join(os.path.dirname(os.path.abspath(__file__)),
                           "ReciApp_Proyecto_Intermodular_DAM.docx")
doc.save(output_path)
print(f"Documento generado: {output_path}")
print(f"Total de figuras: {fig_counter}")
