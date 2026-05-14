#!/usr/bin/env python3
"""
Genera el documento DOCX completo del proyecto ReciApp.
"""

from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml
import os

# ═══════════════════════════════════════════════════════════════
# Helpers
# ═══════════════════════════════════════════════════════════════

GREEN = RGBColor(0x16, 0xA3, 0x4A)
DARK_GREEN = RGBColor(0x15, 0x80, 0x3D)
LIGHT_GREEN = RGBColor(0xDC, 0xFA, 0xDE)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)
BLACK = RGBColor(0x00, 0x00, 0x00)
GRAY = RGBColor(0x6B, 0x72, 0x80)
DARK = RGBColor(0x1F, 0x29, 0x37)
LIGHT_BG = RGBColor(0xF9, 0xFA, 0xFB)
RED_BORDER = RGBColor(0xDC, 0x26, 0x26)

screenshot_counter = 0

def set_cell_shading(cell, color_hex):
    """Aplica color de fondo a una celda."""
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color_hex}" w:val="clear"/>')
    cell._tc.get_or_add_tcPr().append(shading)

def set_cell_border(cell, **kwargs):
    """Set cell border. Use keyword arguments to specify borders (top, bottom, left, right)."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = parse_xml(f'<w:tcBorders {nsdecls("w")}></w:tcBorders>')
    for edge, val in kwargs.items():
        element = parse_xml(
            f'<w:{edge} {nsdecls("w")} w:val="{val.get("val", "single")}" '
            f'w:sz="{val.get("sz", "4")}" w:space="0" '
            f'w:color="{val.get("color", "000000")}"/>'
        )
        tcBorders.append(element)
    tcPr.append(tcBorders)

def add_styled_heading(doc, text, level=1):
    """Añade un heading con estilo verde."""
    heading = doc.add_heading(text, level=level)
    for run in heading.runs:
        run.font.color.rgb = DARK_GREEN if level == 1 else GREEN
    return heading

def add_body(doc, text, bold=False, italic=False, size=11):
    """Añade un párrafo de cuerpo."""
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.size = Pt(size)
    run.font.color.rgb = DARK
    run.bold = bold
    run.italic = italic
    return p

def add_bullet(doc, text, level=0):
    """Añade un bullet point."""
    p = doc.add_paragraph(style='List Bullet')
    p.clear()
    run = p.add_run(text)
    run.font.size = Pt(11)
    run.font.color.rgb = DARK
    if level > 0:
        p.paragraph_format.left_indent = Cm(1.27 * (level + 1))
    return p

def add_screenshot_placeholder(doc, number, description):
    """Añade un recuadro rojo con borde discontinuo y número."""
    global screenshot_counter
    screenshot_counter += 1

    # Crear tabla de 1 celda para el recuadro
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = table.cell(0, 0)

    # Configurar borde rojo discontinuo
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = parse_xml(
        f'<w:tcBorders {nsdecls("w")}>'
        f'  <w:top w:val="dashed" w:sz="12" w:space="0" w:color="DC2626"/>'
        f'  <w:left w:val="dashed" w:sz="12" w:space="0" w:color="DC2626"/>'
        f'  <w:bottom w:val="dashed" w:sz="12" w:space="0" w:color="DC2626"/>'
        f'  <w:right w:val="dashed" w:sz="12" w:space="0" w:color="DC2626"/>'
        f'</w:tcBorders>'
    )
    tcPr.append(tcBorders)

    # Contenido del recuadro
    # Número
    p_num = cell.paragraphs[0]
    p_num.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_num = p_num.add_run(f"Captura #{screenshot_counter}")
    run_num.font.size = Pt(16)
    run_num.font.color.rgb = RED_BORDER
    run_num.bold = True

    p_num.paragraph_format.space_before = Pt(20)

    # Descripción
    p_desc = cell.add_paragraph()
    p_desc.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_desc = p_desc.add_run(description)
    run_desc.font.size = Pt(11)
    run_desc.font.color.rgb = RGBColor(0x99, 0x33, 0x33)
    run_desc.italic = True
    p_desc.paragraph_format.space_after = Pt(20)

    # Espacio interior para que tenga altura
    p_space = cell.add_paragraph()
    p_space.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_space = p_space.add_run("\n\n")
    run_space.font.size = Pt(24)
    p_space.paragraph_format.space_after = Pt(10)

    # Separador después
    doc.add_paragraph()
    return screenshot_counter

def add_styled_table(doc, headers, rows, col_widths=None):
    """Crea una tabla con estilo."""
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = 'Table Grid'

    # Header
    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        set_cell_shading(cell, "158035")
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(header)
        run.font.size = Pt(10)
        run.font.color.rgb = WHITE
        run.bold = True
        run.font.name = 'Calibri'

    # Rows
    for r_idx, row in enumerate(rows):
        for c_idx, val in enumerate(row):
            cell = table.rows[r_idx + 1].cells[c_idx]
            if r_idx % 2 == 1:
                set_cell_shading(cell, "F0FDF4")
            p = cell.paragraphs[0]
            run = p.add_run(str(val))
            run.font.size = Pt(10)
            run.font.color.rgb = DARK
            run.font.name = 'Calibri'

    if col_widths:
        for i, w in enumerate(col_widths):
            for row in table.rows:
                row.cells[i].width = Cm(w)

    doc.add_paragraph()
    return table

def add_code_block(doc, code_text):
    """Simula un bloque de código."""
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    cell = table.cell(0, 0)
    set_cell_shading(cell, "F3F4F6")

    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = parse_xml(
        f'<w:tcBorders {nsdecls("w")}>'
        f'  <w:top w:val="single" w:sz="4" w:space="0" w:color="D1D5DB"/>'
        f'  <w:left w:val="single" w:sz="12" w:space="0" w:color="16A34A"/>'
        f'  <w:bottom w:val="single" w:sz="4" w:space="0" w:color="D1D5DB"/>'
        f'  <w:right w:val="single" w:sz="4" w:space="0" w:color="D1D5DB"/>'
        f'</w:tcBorders>'
    )
    tcPr.append(tcBorders)

    p = cell.paragraphs[0]
    run = p.add_run(code_text)
    run.font.size = Pt(9)
    run.font.name = 'Consolas'
    run.font.color.rgb = DARK
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(6)

    doc.add_paragraph()


# ═══════════════════════════════════════════════════════════════
# DOCUMENTO
# ═══════════════════════════════════════════════════════════════

doc = Document()

# Estilos base
style = doc.styles['Normal']
style.font.name = 'Calibri'
style.font.size = Pt(11)
style.font.color.rgb = DARK

# Márgenes
for section in doc.sections:
    section.top_margin = Cm(2)
    section.bottom_margin = Cm(2)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)

# ═══════════════════════════════════════════════════════════════
# PORTADA
# ═══════════════════════════════════════════════════════════════

# Espacio superior
for _ in range(4):
    doc.add_paragraph()

# Línea decorativa superior
p_line = doc.add_paragraph()
p_line.alignment = WD_ALIGN_PARAGRAPH.CENTER
run_line = p_line.add_run("━" * 50)
run_line.font.color.rgb = GREEN
run_line.font.size = Pt(14)

# Icono
p_icon = doc.add_paragraph()
p_icon.alignment = WD_ALIGN_PARAGRAPH.CENTER
run_icon = p_icon.add_run("♻")
run_icon.font.size = Pt(60)
run_icon.font.color.rgb = GREEN

# Título
p_title = doc.add_paragraph()
p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
run_title = p_title.add_run("ReciApp")
run_title.font.size = Pt(42)
run_title.font.color.rgb = DARK_GREEN
run_title.bold = True

# Subtítulo
p_sub = doc.add_paragraph()
p_sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
run_sub = p_sub.add_run("Sistema de Gestión de Reciclaje y Reducción de Emisiones CO₂")
run_sub.font.size = Pt(16)
run_sub.font.color.rgb = GRAY
run_sub.italic = True

# Línea decorativa inferior
p_line2 = doc.add_paragraph()
p_line2.alignment = WD_ALIGN_PARAGRAPH.CENTER
run_line2 = p_line2.add_run("━" * 50)
run_line2.font.color.rgb = GREEN
run_line2.font.size = Pt(14)

# Información del documento
for _ in range(3):
    doc.add_paragraph()

info_data = [
    ("Documento:", "Documentación Técnica del Proyecto"),
    ("Tipo:", "Proyecto Educativo / TFG"),
    ("Plataformas:", "API REST · App Móvil · App Escritorio · Terminal"),
    ("Fecha:", "Mayo 2026"),
]

for label, value in info_data:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r1 = p.add_run(label + "  ")
    r1.bold = True
    r1.font.size = Pt(12)
    r1.font.color.rgb = DARK_GREEN
    r2 = p.add_run(value)
    r2.font.size = Pt(12)
    r2.font.color.rgb = DARK

# Salto de página
doc.add_page_break()


# ═══════════════════════════════════════════════════════════════
# ÍNDICE
# ═══════════════════════════════════════════════════════════════

add_styled_heading(doc, "Índice de Contenidos", level=1)

indice = [
    "1. Resumen del Proyecto",
    "2. Tecnologías Utilizadas",
    "3. Arquitectura del Sistema",
    "4. Modelo Entidad-Relación",
    "5. Estructura de la Base de Datos",
    "6. Roles y Permisos",
    "7. API REST (Spring Boot)",
    "8. Aplicación Móvil (React Native / Expo)",
    "9. Aplicación de Escritorio (JavaFX)",
    "10. Terminal de Reciclaje (Python)",
    "11. Despliegue y Dockerización",
    "12. Casos de Uso",
    "13. Capturas de Pantalla",
]

for item in indice:
    p = doc.add_paragraph()
    run = p.add_run(item)
    run.font.size = Pt(12)
    run.font.color.rgb = DARK
    p.paragraph_format.space_after = Pt(4)

doc.add_page_break()


# ═══════════════════════════════════════════════════════════════
# 1. RESUMEN DEL PROYECTO
# ═══════════════════════════════════════════════════════════════

add_styled_heading(doc, "1. Resumen del Proyecto", level=1)

add_body(doc,
    "ReciApp es un sistema integral de gestión de reciclaje diseñado para "
    "fomentar la reducción de emisiones de CO₂ mediante el registro y seguimiento "
    "de materiales reciclados. El sistema permite a los usuarios registrar productos "
    "reciclables, consultar su historial de reciclaje y visualizar su impacto "
    "ambiental medido en kilogramos de CO₂ reducidos.")

add_body(doc,
    "El proyecto consta de cuatro componentes principales que se comunican "
    "a través de una API REST centralizada:")

add_bullet(doc, "API REST (Spring Boot): Servidor central que gestiona toda la lógica de negocio, autenticación JWT y persistencia en base de datos MySQL.")
add_bullet(doc, "Aplicación Móvil (React Native / Expo): Cliente multiplataforma para que los usuarios consulten productos, vean su historial de reciclaje y gestionen su perfil.")
add_bullet(doc, "Aplicación de Escritorio (JavaFX): Herramienta de administración con interfaz gráfica completa para gestión de usuarios, productos y transacciones.")
add_bullet(doc, "Terminal de Reciclaje (Python): Script de terminal para puntos de reciclaje físicos, donde se escanean códigos de barras y se identifican usuarios mediante su TAP (Token de Autenticación Personal).")

add_body(doc, "")
add_body(doc,
    "Cada producto registrado en el sistema tiene asociado un valor de emisiones "
    "reducibles (en kg de CO₂). Cuando un usuario recicla un producto, ese valor "
    "se acumula en su perfil, permitiéndole visualizar su contribución personal "
    "a la reducción de emisiones.",
    size=11)

add_screenshot_placeholder(doc, 1, "Diagrama general de la arquitectura del sistema (API central + 3 clientes)")

doc.add_page_break()


# ═══════════════════════════════════════════════════════════════
# 2. TECNOLOGÍAS UTILIZADAS
# ═══════════════════════════════════════════════════════════════

add_styled_heading(doc, "2. Tecnologías Utilizadas", level=1)

add_styled_heading(doc, "2.1 Backend — API REST", level=2)

add_styled_table(doc,
    ["Tecnología", "Versión", "Propósito"],
    [
        ["Java", "21 (LTS)", "Lenguaje principal del backend"],
        ["Spring Boot", "3.4.4", "Framework web y contenedor de aplicaciones"],
        ["Spring Security", "6.x", "Autenticación y autorización con JWT"],
        ["Spring Data JPA", "3.x", "ORM y acceso a base de datos"],
        ["Hibernate", "6.x", "Implementación JPA para mapeo objeto-relacional"],
        ["JJWT", "0.12.6", "Generación y validación de tokens JWT"],
        ["Lombok", "1.18.x", "Reducción de código boilerplate"],
        ["Gradle", "8.x", "Herramienta de construcción y gestión de dependencias"],
        ["MySQL", "8.x", "Sistema de gestión de base de datos relacional"],
        ["BCrypt", "—", "Algoritmo de hash para contraseñas (strength 10)"],
    ],
    col_widths=[4, 2.5, 9]
)

add_styled_heading(doc, "2.2 Aplicación Móvil", level=2)

add_styled_table(doc,
    ["Tecnología", "Versión", "Propósito"],
    [
        ["React Native", "0.81.5", "Framework de desarrollo móvil multiplataforma"],
        ["Expo SDK", "54", "Plataforma de desarrollo y despliegue simplificado"],
        ["Expo Router", "4.x", "Navegación basada en sistema de archivos"],
        ["TypeScript", "5.x", "Tipado estático sobre JavaScript"],
        ["NativeWind", "4.x", "Tailwind CSS adaptado para React Native"],
        ["Zustand", "5.x", "Gestión de estado global ligera y reactiva"],
        ["React Native Paper", "5.x", "Librería de componentes Material Design"],
        ["AsyncStorage", "2.x", "Almacenamiento local persistente clave-valor"],
        ["Ionicons", "—", "Iconos vectoriales para la interfaz"],
    ],
    col_widths=[4, 2.5, 9]
)

add_styled_heading(doc, "2.3 Aplicación de Escritorio", level=2)

add_styled_table(doc,
    ["Tecnología", "Versión", "Propósito"],
    [
        ["Java", "24", "Lenguaje principal de la aplicación de escritorio"],
        ["JavaFX", "25", "Framework de interfaz gráfica de usuario"],
        ["FXML", "—", "Formato declarativo para definir interfaces gráficas"],
        ["Gson", "2.x", "Serialización/deserialización JSON"],
        ["HttpClient", "Java 11+", "Cliente HTTP nativo de Java para llamadas a la API"],
        ["ChartFX", "—", "Librería de gráficos para visualización de datos"],
        ["CSS", "—", "Estilización personalizada de componentes JavaFX"],
    ],
    col_widths=[4, 2.5, 9]
)

add_styled_heading(doc, "2.4 Terminal de Reciclaje", level=2)

add_styled_table(doc,
    ["Tecnología", "Versión", "Propósito"],
    [
        ["Python", "3.10+", "Lenguaje de scripting para el terminal"],
        ["requests", "2.x", "Librería HTTP para comunicación con la API"],
        ["ANSI Colors", "—", "Colorización de la salida en terminal"],
    ],
    col_widths=[4, 2.5, 9]
)

add_styled_heading(doc, "2.5 Infraestructura y Despliegue", level=2)

add_styled_table(doc,
    ["Tecnología", "Versión", "Propósito"],
    [
        ["AWS EC2", "—", "Servidor en la nube para alojar la API y la base de datos"],
        ["Ubuntu Server", "22.04 LTS", "Sistema operativo del servidor"],
        ["Docker", "24.x", "Contenedorización de la API y MySQL"],
        ["Docker Compose", "2.x", "Orquestación de contenedores multi-servicio"],
        ["Nginx", "—", "Proxy inverso (opcional, preparado para producción)"],
    ],
    col_widths=[4, 2.5, 9]
)

doc.add_page_break()


# ═══════════════════════════════════════════════════════════════
# 3. ARQUITECTURA DEL SISTEMA
# ═══════════════════════════════════════════════════════════════

add_styled_heading(doc, "3. Arquitectura del Sistema", level=1)

add_body(doc,
    "ReciApp sigue una arquitectura cliente-servidor donde una API REST "
    "centralizada actúa como único punto de acceso a la base de datos. "
    "Los tres clientes (móvil, escritorio y terminal) se comunican con la API "
    "mediante peticiones HTTP, utilizando JSON como formato de intercambio de datos.")

add_styled_heading(doc, "3.1 Diagrama de Arquitectura", level=2)

add_body(doc,
    "La comunicación entre todos los componentes se realiza a través de HTTP "
    "en el puerto 3000. Se evaluó el uso de HTTPS con certificados autofirmados, "
    "pero se descartó por incompatibilidad con el entorno de desarrollo Expo Go "
    "de React Native, que no permite personalizar la configuración de seguridad "
    "de red de Android sin eyectar del flujo gestionado.",
    size=11)

add_code_block(doc,
"""┌─────────────────┐     HTTP :3000      ┌──────────────────┐     JDBC      ┌──────────┐
│   App Móvil     │ ──────────────────> │                  │ ──────────> │          │
│  (React Native) │                     │   API REST       │             │  MySQL   │
├─────────────────┤     HTTP :3000      │  (Spring Boot)   │             │  8.x     │
│  App Escritorio │ ──────────────────> │                  │             │          │
│    (JavaFX)     │                     │  Puerto: 3000    │             │ Puerto:  │
├─────────────────┤     HTTP :3000      │  Host: 0.0.0.0   │             │  3306    │
│   Terminal      │ ──────────────────> │                  │             │          │
│   (Python)      │                     └──────────────────┘             └──────────┘
└─────────────────┘""")

add_styled_heading(doc, "3.2 Flujo de Autenticación", level=2)

add_body(doc,
    "El sistema utiliza JSON Web Tokens (JWT) para la autenticación. "
    "El flujo es el siguiente:")

add_bullet(doc, "El cliente envía credenciales (nombre + contraseña) a POST /api/usuarios/login.")
add_bullet(doc, "El servidor valida las credenciales con BCrypt y genera un token JWT con claims: userId, nombre y permisos.")
add_bullet(doc, "El token tiene una validez de 7 días (604.800.000 ms).")
add_bullet(doc, "En cada petición posterior, el cliente envía el token en la cabecera Authorization: Bearer <token>.")
add_bullet(doc, "El filtro JwtAuthFilter intercepta la petición, valida el token y establece el contexto de seguridad de Spring.")

add_screenshot_placeholder(doc, 2, "Flujo de autenticación JWT: diagrama de secuencia login → token → peticiones autenticadas")

doc.add_page_break()


# ═══════════════════════════════════════════════════════════════
# 4. MODELO ENTIDAD-RELACIÓN
# ═══════════════════════════════════════════════════════════════

add_styled_heading(doc, "4. Modelo Entidad-Relación", level=1)

add_body(doc,
    "La base de datos reciInventario_db consta de tres tablas principales "
    "con las siguientes relaciones:")

add_code_block(doc,
"""    ┌──────────────────┐           ┌──────────────────┐
    │     USUARIOS     │           │    PRODUCTOS     │
    ├──────────────────┤           ├──────────────────┤
    │ PK Id_Usuario    │           │ PK Tipo          │
    │    Nombre (UQ)   │           │ PK Numero_barras │
    │    Hash_Contraseña│          │    Nombre        │
    │    Permisos      │           │    Emisiones_Red.│
    │    Emisiones_Red.│           │    Material      │
    │    TAP           │           │    Imagen        │
    └──────┬───────────┘           └──────┬───────────┘
           │ 1                            │ 1
           │                              │
           │ N                            │ N
    ┌──────┴──────────────────────────────┴───────────┐
    │                   RECICLA                       │
    ├─────────────────────────────────────────────────┤
    │ PK,FK Id_Usuario                                │
    │ PK,FK Tipo                                      │
    │ PK,FK Numero_barras                             │
    │ PK    Fecha                                     │
    │ PK    Hora                                      │
    └─────────────────────────────────────────────────┘""")

add_body(doc, "Relaciones:", bold=True)
add_bullet(doc, "Un Usuario puede tener muchas entradas en Recicla (1:N).")
add_bullet(doc, "Un Producto puede aparecer en muchas entradas en Recicla (1:N).")
add_bullet(doc, "La tabla Recicla es la tabla intermedia que registra cada acción de reciclaje con fecha y hora exactas.")
add_bullet(doc, "Ambas relaciones tienen ON DELETE CASCADE: si se elimina un usuario o producto, se eliminan sus registros de reciclaje asociados.")

add_screenshot_placeholder(doc, 3, "Modelo Entidad-Relación generado desde MySQL Workbench o herramienta similar")

doc.add_page_break()


# ═══════════════════════════════════════════════════════════════
# 5. ESTRUCTURA DE LA BASE DE DATOS
# ═══════════════════════════════════════════════════════════════

add_styled_heading(doc, "5. Estructura de la Base de Datos", level=1)

add_body(doc,
    "La base de datos utiliza MySQL 8.x con el motor InnoDB y codificación "
    "utf8mb4 (soporte completo de Unicode). A continuación se detalla cada tabla.")

add_styled_heading(doc, "5.1 Tabla: Usuarios", level=2)

add_body(doc, "Almacena la información de todos los usuarios del sistema, tanto clientes como administradores.")

add_styled_table(doc,
    ["Columna", "Tipo", "Restricciones", "Descripción"],
    [
        ["Id_Usuario", "INT", "PK, AUTO_INCREMENT", "Identificador único del usuario"],
        ["Nombre", "VARCHAR(50)", "NOT NULL, UNIQUE", "Nombre de usuario para login"],
        ["Hash_Contraseña", "VARCHAR(100)", "NOT NULL", "Hash BCrypt de la contraseña"],
        ["Permisos", "VARCHAR(15)", "DEFAULT 'cliente'", "Rol: 'cliente' o 'administrador'"],
        ["Emisiones_Reducidas", "FLOAT", "DEFAULT 0", "kg de CO₂ reducidos acumulados"],
        ["TAP", "INT", "DEFAULT NULL", "Token de Autenticación Personal (6 dígitos)"],
    ],
    col_widths=[3.5, 3, 3.5, 5.5]
)

add_styled_heading(doc, "5.2 Tabla: Productos", level=2)

add_body(doc, "Catálogo de productos reciclables. Utiliza una clave primaria compuesta por el tipo de código y el número de barras.")

add_styled_table(doc,
    ["Columna", "Tipo", "Restricciones", "Descripción"],
    [
        ["Tipo", "VARCHAR(10)", "PK (compuesta)", "Tipo de código (p.ej. 'EAN13')"],
        ["Numero_barras", "BIGINT", "PK (compuesta)", "Número del código de barras"],
        ["Nombre", "VARCHAR(50)", "—", "Nombre comercial del producto"],
        ["Emisiones_Reducibles", "FLOAT", "—", "kg de CO₂ que se reducen al reciclar"],
        ["Material", "VARCHAR(15)", "—", "Material principal (PET, Vidrio, Aluminio…)"],
        ["Imagen", "LONGTEXT", "—", "Imagen codificada en Base64"],
    ],
    col_widths=[3.5, 3, 3.5, 5.5]
)

add_styled_heading(doc, "5.3 Tabla: Recicla", level=2)

add_body(doc, "Registra cada acción de reciclaje. Su clave primaria compuesta de cinco columnas permite registrar múltiples reciclajes del mismo producto por el mismo usuario en momentos diferentes.")

add_styled_table(doc,
    ["Columna", "Tipo", "Restricciones", "Descripción"],
    [
        ["Id_Usuario", "INT", "PK, FK → Usuarios", "Usuario que realizó el reciclaje"],
        ["Tipo", "VARCHAR(10)", "PK, FK → Productos", "Tipo de código del producto"],
        ["Numero_barras", "BIGINT", "PK, FK → Productos", "Código de barras del producto"],
        ["Fecha", "DATE", "PK", "Fecha del reciclaje"],
        ["Hora", "TIME", "PK", "Hora exacta del reciclaje"],
    ],
    col_widths=[3.5, 3, 3.5, 5.5]
)

add_styled_heading(doc, "5.4 Datos de Ejemplo", level=2)

add_body(doc, "El script de inicialización incluye cinco productos de ejemplo para pruebas:")

add_styled_table(doc,
    ["Producto", "Código EAN13", "Material", "CO₂ (kg)"],
    [
        ["Coca-Cola 330ml", "8410076472885", "Aluminio", "0.033"],
        ["Agua Bezoya 1.5L", "8411700000017", "PET", "0.025"],
        ["Leche Hacendado 1L", "8480000291455", "Brick", "0.040"],
        ["Mahou Clásica 330ml", "8410128800505", "Vidrio", "0.045"],
        ["Fanta Naranja 330ml", "5449000000996", "Aluminio", "0.033"],
    ],
    col_widths=[4.5, 3.5, 3, 2.5]
)

doc.add_page_break()


# ═══════════════════════════════════════════════════════════════
# 6. ROLES Y PERMISOS
# ═══════════════════════════════════════════════════════════════

add_styled_heading(doc, "6. Roles y Permisos", level=1)

add_body(doc,
    "El sistema define dos roles de usuario almacenados en la columna Permisos "
    "de la tabla Usuarios. La autorización se implementa tanto a nivel de API "
    "(Spring Security) como a nivel de interfaz (visibilidad condicional de opciones).")

add_styled_heading(doc, "6.1 Rol: Cliente", level=2)

add_body(doc, "Es el rol por defecto asignado a todo usuario nuevo al registrarse. Permite:")

add_bullet(doc, "Iniciar sesión y registrarse en cualquiera de las plataformas.")
add_bullet(doc, "Consultar el catálogo completo de productos reciclables.")
add_bullet(doc, "Visualizar su historial personal de reciclaje.")
add_bullet(doc, "Ver su perfil con las emisiones de CO₂ reducidas acumuladas.")
add_bullet(doc, "Solicitar y renovar su TAP (Token de Autenticación Personal).")
add_bullet(doc, "Modificar su nombre de usuario (requiere contraseña actual).")
add_bullet(doc, "Cambiar su contraseña.")
add_bullet(doc, "Eliminar su cuenta (requiere contraseña actual).")
add_bullet(doc, "Reciclar productos en terminales físicos usando su TAP.")

add_styled_heading(doc, "6.2 Rol: Administrador", level=2)

add_body(doc, "Tiene todas las capacidades del cliente, más las siguientes funciones de administración:")

add_bullet(doc, "Gestión completa de Usuarios: crear, listar, modificar y eliminar cualquier usuario.")
add_bullet(doc, "Gestión completa de Productos: crear, listar, modificar y eliminar productos del catálogo, incluyendo imágenes en Base64.")
add_bullet(doc, "Gestión completa de Transacciones: crear, listar, modificar y eliminar registros de reciclaje.")
add_bullet(doc, "Visualizar gráficos de emisiones reducidas de todos los usuarios (app escritorio).")
add_bullet(doc, "Operar el terminal de reciclaje en puntos físicos.")

add_styled_heading(doc, "6.3 Modo Invitado (Solo App Móvil)", level=2)

add_body(doc,
    "La aplicación móvil ofrece un modo invitado para uso sin conexión o sin registro. "
    "El usuario invitado puede explorar el catálogo de productos con datos de ejemplo, "
    "pero no puede ver su historial real, solicitar TAP ni gestionar su perfil. "
    "Se muestra un banner indicativo en las pantallas que requieren autenticación.")

add_styled_heading(doc, "6.4 Matriz de Permisos", level=2)

add_styled_table(doc,
    ["Acción", "Cliente", "Admin", "Invitado"],
    [
        ["Registro / Login", "✅", "✅", "❌"],
        ["Ver productos", "✅", "✅", "✅ (ejemplo)"],
        ["Ver historial propio", "✅", "✅", "❌"],
        ["Ver/solicitar TAP", "✅", "✅", "❌"],
        ["Modificar nombre/contraseña", "✅", "✅", "❌"],
        ["Eliminar cuenta propia", "✅", "✅", "❌"],
        ["CRUD Usuarios (admin)", "❌", "✅", "❌"],
        ["CRUD Productos (admin)", "❌", "✅", "❌"],
        ["CRUD Transacciones (admin)", "❌", "✅", "❌"],
        ["Operar terminal", "❌", "✅", "❌"],
        ["Gráficos de emisiones", "❌", "✅", "❌"],
    ],
    col_widths=[6, 2.5, 2.5, 2.5]
)

doc.add_page_break()


# ═══════════════════════════════════════════════════════════════
# 7. API REST (Spring Boot)
# ═══════════════════════════════════════════════════════════════

add_styled_heading(doc, "7. API REST (Spring Boot)", level=1)

add_body(doc,
    "La API REST es el componente central del sistema. Desarrollada con "
    "Spring Boot 3.4.4 y Java 21, expone todos los endpoints necesarios "
    "para la gestión de usuarios, productos e historial de reciclaje. "
    "Escucha en el puerto 3000 (HTTP) y se comunica con MySQL en el puerto 3306.")

add_styled_heading(doc, "7.1 Estructura del Proyecto", level=2)

add_code_block(doc,
"""api-spring/src/main/java/com/reciapp/api/
├── config/
│   └── SecurityConfig.java        ← Configuración de Spring Security
├── controller/
│   ├── AdminController.java       ← Endpoints de administración
│   ├── HealthController.java      ← Endpoint de salud
│   ├── HistorialController.java   ← Gestión de historial
│   ├── ProductoController.java    ← Gestión de productos
│   ├── UsuarioController.java     ← Auth + perfil de usuario
│   └── GlobalExceptionHandler.java
├── dto/                           ← Objetos de transferencia de datos
│   ├── LoginRequest.java, RegisterRequest.java
│   ├── AuthResponse.java, UsuarioDto.java
│   ├── ProductoDto.java, HistorialDto.java
│   ├── Admin*.java                ← DTOs específicos de admin
│   └── ...
├── entity/                        ← Entidades JPA
│   ├── Usuario.java
│   ├── Producto.java, ProductoId.java
│   └── Recicla.java, ReciclaId.java
├── repository/                    ← Repositorios Spring Data JPA
│   ├── UsuarioRepository.java
│   ├── ProductoRepository.java
│   └── ReciclaRepository.java
├── security/                      ← Seguridad JWT
│   ├── JwtService.java
│   └── JwtAuthFilter.java
└── service/                       ← Lógica de negocio
    ├── UsuarioService.java
    ├── ProductoService.java
    ├── ReciclaService.java
    └── AdminService.java""")

add_styled_heading(doc, "7.2 Endpoints de la API", level=2)

add_body(doc, "Endpoints Públicos (sin autenticación):", bold=True)

add_styled_table(doc,
    ["Método", "Endpoint", "Descripción"],
    [
        ["GET", "/api/health", "Comprobación de estado del servidor"],
        ["POST", "/api/usuarios/login", "Autenticación (devuelve JWT)"],
        ["POST", "/api/usuarios/register", "Registro de nuevo usuario"],
    ],
    col_widths=[2, 5.5, 7]
)

add_body(doc, "Endpoints Protegidos (requieren JWT):", bold=True)

add_styled_table(doc,
    ["Método", "Endpoint", "Descripción"],
    [
        ["GET", "/api/usuarios/profile", "Obtener perfil del usuario autenticado"],
        ["PUT", "/api/usuarios/{id}/nombre", "Actualizar nombre de usuario"],
        ["PUT", "/api/usuarios/{id}/password", "Cambiar contraseña"],
        ["PUT", "/api/usuarios/{id}/tap", "Solicitar/renovar TAP"],
        ["DELETE", "/api/usuarios/{id}", "Eliminar cuenta propia"],
        ["GET", "/api/usuarios/by-tap/{tap}", "Buscar usuario por TAP"],
        ["GET", "/api/productos", "Listar todos los productos"],
        ["GET", "/api/productos/barcode/{code}", "Buscar producto por código de barras"],
        ["GET", "/api/historial", "Obtener historial del usuario"],
        ["POST", "/api/historial", "Registrar nuevo reciclaje"],
    ],
    col_widths=[2, 5.5, 7]
)

add_body(doc, "Endpoints de Administración (requieren JWT + rol administrador):", bold=True)

add_styled_table(doc,
    ["Método", "Endpoint", "Descripción"],
    [
        ["GET", "/api/admin/usuarios", "Listar todos los usuarios"],
        ["POST", "/api/admin/usuarios", "Crear usuario (admin)"],
        ["PUT", "/api/admin/usuarios/{id}", "Modificar usuario"],
        ["DELETE", "/api/admin/usuarios/{id}", "Eliminar usuario"],
        ["GET", "/api/admin/productos", "Listar todos los productos"],
        ["POST", "/api/admin/productos", "Crear producto"],
        ["PUT", "/api/admin/productos/{tipo}/{barras}", "Modificar producto"],
        ["DELETE", "/api/admin/productos/{tipo}/{barras}", "Eliminar producto"],
        ["GET", "/api/admin/transacciones", "Listar todas las transacciones"],
        ["POST", "/api/admin/transacciones", "Crear transacción"],
        ["DELETE", "/api/admin/transacciones", "Eliminar transacción"],
    ],
    col_widths=[2, 5.5, 7]
)

add_styled_heading(doc, "7.3 Seguridad y Autenticación", level=2)

add_body(doc,
    "La seguridad se implementa con Spring Security 6 en modo stateless. "
    "Las contraseñas se almacenan hasheadas con BCrypt (strength 10). "
    "Los tokens JWT incluyen tres claims personalizados: userId, nombre y permisos. "
    "El filtro JwtAuthFilter se ejecuta antes del UsernamePasswordAuthenticationFilter "
    "en la cadena de filtros de seguridad.")

add_body(doc,
    "CORS está habilitado para todos los orígenes (*) para facilitar el "
    "desarrollo. CSRF está deshabilitado dado que la API es stateless y "
    "utiliza tokens en la cabecera Authorization en lugar de cookies de sesión.")

add_styled_heading(doc, "7.4 Generación del TAP", level=2)

add_body(doc,
    "El TAP (Token de Autenticación Personal) es un número aleatorio de 6 dígitos "
    "(rango 100000-999999) único para cada usuario. Se utiliza en los terminales "
    "de reciclaje para identificar al usuario sin necesidad de contraseña. "
    "El sistema intenta generar un TAP único hasta 20 veces antes de fallar.")

add_screenshot_placeholder(doc, 4, "Captura de Postman o similar mostrando una petición a la API con respuesta JSON")

doc.add_page_break()


# ═══════════════════════════════════════════════════════════════
# 8. APLICACIÓN MÓVIL
# ═══════════════════════════════════════════════════════════════

add_styled_heading(doc, "8. Aplicación Móvil (React Native / Expo)", level=1)

add_body(doc,
    "La aplicación móvil está desarrollada con React Native 0.81.5 y Expo SDK 54, "
    "utilizando TypeScript como lenguaje principal. La navegación se basa en "
    "Expo Router (file-based routing) y el estilizado utiliza NativeWind "
    "(adaptación de Tailwind CSS para React Native).")

add_styled_heading(doc, "8.1 Estructura de Navegación", level=2)

add_code_block(doc,
"""app/
├── _layout.tsx              ← Layout raíz: auth guard + conexión
├── index.tsx                ← Punto de entrada
├── connection-modal.tsx     ← Modal de conexión offline
├── (auth)/                  ← Grupo de autenticación
│   ├── _layout.tsx
│   ├── login.tsx            ← Pantalla de inicio de sesión
│   └── register.tsx         ← Pantalla de registro
└── (tabs)/                  ← Navegación por pestañas
    ├── _layout.tsx          ← Tab navigator (3 pestañas)
    ├── productos.tsx        ← Catálogo de productos
    ├── productos/[barcode].tsx  ← Detalle de producto
    ├── historial.tsx        ← Historial de reciclaje
    └── perfil.tsx           ← Perfil del usuario""")

add_styled_heading(doc, "8.2 Gestión de Estado (Zustand)", level=2)

add_body(doc,
    "La aplicación utiliza Zustand para la gestión de estado global, "
    "con cuatro stores independientes:")

add_styled_table(doc,
    ["Store", "Responsabilidad", "Persistencia"],
    [
        ["authStore", "Token JWT, datos de usuario, login/logout, modo invitado", "AsyncStorage"],
        ["connectionStore", "Estado de conexión con la API, flag de verificación", "No"],
        ["productStore", "Caché de la lista de productos", "AsyncStorage"],
        ["recycleStore", "Historial de reciclaje del usuario", "AsyncStorage"],
    ],
    col_widths=[3.5, 7, 3.5]
)

add_styled_heading(doc, "8.3 Flujo de Inicio de la App", level=2)

add_body(doc,
    "Al iniciar la aplicación, el layout raíz ejecuta tres procesos en orden:")

add_bullet(doc, "1. Restaurar sesión: Se intenta cargar un token JWT y datos de usuario desde AsyncStorage.")
add_bullet(doc, "2. Comprobar conexión: Se hace una petición GET a /api/health. Si falla, se muestra el modal de conexión con opciones de recargar o entrar en modo offline.")
add_bullet(doc, "3. Auth guard: Una vez comprobada la conexión, si no hay token se redirige a login; si hay token y el usuario está en la pantalla de auth, se redirige a productos.")

add_body(doc,
    "Este flujo resuelve una race condition detectada durante el desarrollo: "
    "el auth guard no debe actuar hasta que la comprobación de conexión haya "
    "finalizado (controlado por el flag hasChecked).",
    italic=True, size=10)

add_styled_heading(doc, "8.4 Pantallas Principales", level=2)

add_body(doc, "Login y Registro", bold=True)
add_body(doc,
    "Formularios con validación de campos obligatorios. El login envía nombre "
    "y contraseña a la API y almacena el token JWT en AsyncStorage. "
    "Incluye enlace para navegar al registro y viceversa.")

add_body(doc, "Productos", bold=True)
add_body(doc,
    "Lista todos los productos de la base de datos en tarjetas (Cards) con "
    "información del nombre, material, código de barras y emisiones reducibles. "
    "Incluye barra de búsqueda por nombre/código y filtro por material "
    "(Todos, PET, Vidrio, Aluminio, Plástico, Brick). Soporta pull-to-refresh.")

add_body(doc, "Historial", bold=True)
add_body(doc,
    "Muestra el historial personal de reciclaje con un resumen del total de "
    "CO₂ reducido en la cabecera. Cada entrada muestra el producto reciclado, "
    "fecha y emisiones. En modo invitado muestra datos de ejemplo con un banner "
    "de advertencia.")

add_body(doc, "Perfil", bold=True)
add_body(doc,
    "Panel completo de gestión del usuario: estadísticas de emisiones (en kg y "
    "toneladas), modal para ver/solicitar/renovar TAP, formularios para cambiar "
    "nombre y contraseña, y opción de eliminar cuenta. En modo invitado, muestra "
    "un botón para iniciar sesión.")

add_styled_heading(doc, "8.5 Servicio API y Modo Offline", level=2)

add_body(doc,
    "La comunicación con la API se abstrae mediante una interfaz ApiService "
    "con dos implementaciones: RealApiService (peticiones HTTP reales) y "
    "OfflineApiService (datos simulados para modo sin conexión). "
    "Una factoría selecciona la implementación según el estado de conexión.")

add_screenshot_placeholder(doc, 5, "Pantalla de login de la app móvil")
add_screenshot_placeholder(doc, 6, "Pantalla de productos con filtro de materiales activo")
add_screenshot_placeholder(doc, 7, "Pantalla de historial mostrando reciclajes y total CO₂")
add_screenshot_placeholder(doc, 8, "Pantalla de perfil con estadísticas y TAP")
add_screenshot_placeholder(doc, 9, "Modal de conexión offline")

doc.add_page_break()


# ═══════════════════════════════════════════════════════════════
# 9. APLICACIÓN DE ESCRITORIO (JavaFX)
# ═══════════════════════════════════════════════════════════════

add_styled_heading(doc, "9. Aplicación de Escritorio (JavaFX)", level=1)

add_body(doc,
    "La aplicación de escritorio está desarrollada con JavaFX 25 y Java 24. "
    "Proporciona una interfaz gráfica completa para la administración del sistema, "
    "incluyendo la gestión de usuarios, productos y transacciones de reciclaje. "
    "Se comunica con la API REST a través de HttpClient de Java.")

add_styled_heading(doc, "9.1 Estructura de Vistas", level=2)

add_styled_table(doc,
    ["Vista (FXML)", "Controlador", "Descripción"],
    [
        ["loginStart_win.fxml", "LoginController", "Pantalla de inicio de sesión"],
        ["singUp_win.fxml", "SingUpController", "Formulario de registro de usuario"],
        ["main_win.fxml", "MainController", "Dashboard principal con 3 pestañas"],
        ["settings_win.fxml", "SettingsController", "Configuración del usuario"],
        ["changePasswd.fxml", "ChangePasswd", "Cambio de contraseña"],
        ["escanear_win.fxml", "Escanear", "Interfaz de escaneo de códigos de barras"],
        ["newProducto.fxml", "NewProducto", "Formulario para crear producto"],
        ["modProducto.fxml", "ModProducto", "Formulario para modificar producto"],
        ["newUser.fxml", "NewUser", "Formulario para crear usuario"],
        ["modUser.fxml", "ModUser", "Formulario para modificar usuario"],
        ["newTransaccion.fxml", "NewTransaccion", "Formulario para crear transacción"],
        ["modTransaccion.fxml", "ModTransaccion", "Formulario para modificar transacción"],
    ],
    col_widths=[4, 3.5, 6.5]
)

add_styled_heading(doc, "9.2 Dashboard Principal (MainController)", level=2)

add_body(doc, "La ventana principal contiene tres pestañas:")

add_body(doc, "Pestaña de Información Personal", bold=True)
add_body(doc,
    "Muestra los datos del usuario actual: nombre, rol y emisiones "
    "de CO₂ reducidas acumuladas. Permite acceder a la configuración "
    "y al cambio de contraseña.")

add_body(doc, "Pestaña de Administración (solo administradores)", bold=True)
add_body(doc,
    "Visible únicamente para usuarios con rol 'administrador'. Contiene "
    "tres sub-secciones con tablas para gestionar usuarios, productos y "
    "transacciones. Cada tabla tiene botones para crear, modificar y "
    "eliminar registros, que abren ventanas de diálogo FXML específicas.")

add_body(doc, "Pestaña de Gráfico de Emisiones", bold=True)
add_body(doc,
    "Muestra un gráfico de barras (usando ChartFX) con las emisiones "
    "reducidas por cada usuario del sistema. Permite visualizar de forma "
    "comparativa el impacto de cada usuario.")

add_styled_heading(doc, "9.3 ApiClient (Patrón Singleton)", level=2)

add_body(doc,
    "Toda la comunicación con la API se centraliza en la clase ApiClient, "
    "que implementa el patrón Singleton. Utiliza HttpClient de Java "
    "(java.net.http) para realizar las peticiones HTTP y Gson para "
    "serializar/deserializar JSON. La configuración se carga desde "
    "un fichero configuration.properties que contiene la URL de la API.")

add_code_block(doc,
"""# configuration.properties
# Configuracion de la API REST
api.url=http://52.201.91.206:3000""")

add_body(doc,
    "Al arrancar, la aplicación comprueba la conectividad con la API "
    "mediante una petición al endpoint /api/health. Si no hay conexión, "
    "muestra un mensaje de error y no permite continuar.")

add_screenshot_placeholder(doc, 10, "Pantalla de login de la aplicación JavaFX")
add_screenshot_placeholder(doc, 11, "Dashboard principal - pestaña de información personal")
add_screenshot_placeholder(doc, 12, "Dashboard principal - pestaña de administración con tabla de usuarios")
add_screenshot_placeholder(doc, 13, "Dashboard principal - pestaña de administración con tabla de productos")
add_screenshot_placeholder(doc, 14, "Dashboard principal - gráfico de emisiones por usuario")
add_screenshot_placeholder(doc, 15, "Ventana de diálogo para crear/modificar un producto")

doc.add_page_break()


# ═══════════════════════════════════════════════════════════════
# 10. TERMINAL DE RECICLAJE (Python)
# ═══════════════════════════════════════════════════════════════

add_styled_heading(doc, "10. Terminal de Reciclaje (Python)", level=1)

add_body(doc,
    "El terminal de reciclaje es un script de Python diseñado para ejecutarse "
    "en puntos de reciclaje físicos (kioscos, contenedores inteligentes). "
    "Utiliza un lector de códigos de barras USB y la identificación del usuario "
    "mediante su TAP para registrar acciones de reciclaje.")

add_styled_heading(doc, "10.1 Flujo de Funcionamiento", level=2)

add_bullet(doc, "1. Verificación de conexión: El terminal comprueba la disponibilidad de la API (GET /api/health).")
add_bullet(doc, "2. Autenticación del operador: Se autentica con credenciales de administrador (configurables por variables de entorno o entrada manual, 3 intentos).")
add_bullet(doc, "3. Bucle principal de reciclaje:")
add_bullet(doc, "   a) Escaneo del producto: Se lee el código de barras y se busca en la API (GET /api/productos/barcode/{codigo}).", level=1)
add_bullet(doc, "   b) Identificación del usuario: Se introduce el TAP del usuario y se verifica en la API (GET /api/usuarios/by-tap/{tap}).", level=1)
add_bullet(doc, "   c) Confirmación: Se muestra un resumen y se pide confirmación.", level=1)
add_bullet(doc, "   d) Registro: Se envía la transacción a la API (POST /api/historial) y se muestra el resultado con las emisiones acumuladas.", level=1)

add_styled_heading(doc, "10.2 Configuración", level=2)

add_styled_table(doc,
    ["Variable de Entorno", "Valor por Defecto", "Descripción"],
    [
        ["RECIAPP_API_URL", "http://52.201.91.206:3000", "URL base de la API REST"],
        ["RECIAPP_ADMIN_USER", "terminal", "Usuario administrador del terminal"],
        ["RECIAPP_ADMIN_PASS", "1234", "Contraseña del operador"],
    ],
    col_widths=[4.5, 5, 5]
)

add_styled_heading(doc, "10.3 Interfaz de Usuario", level=2)

add_body(doc,
    "El terminal utiliza colores ANSI para proporcionar una experiencia visual "
    "clara en la consola: verde para éxito, rojo para errores, amarillo para "
    "advertencias y cian para datos importantes. Los recuadros decorativos "
    "con caracteres Unicode delimitan secciones como el banner de bienvenida "
    "y la confirmación de reciclaje.")

add_screenshot_placeholder(doc, 16, "Terminal de reciclaje: banner de bienvenida y autenticación del operador")
add_screenshot_placeholder(doc, 17, "Terminal de reciclaje: flujo completo de un reciclaje (escaneo → TAP → confirmación)")

doc.add_page_break()


# ═══════════════════════════════════════════════════════════════
# 11. DESPLIEGUE Y DOCKERIZACIÓN
# ═══════════════════════════════════════════════════════════════

add_styled_heading(doc, "11. Despliegue y Dockerización", level=1)

add_styled_heading(doc, "11.1 Infraestructura Actual", level=2)

add_body(doc,
    "El sistema se despliega actualmente en una instancia AWS EC2 con Ubuntu "
    "Server. La API y la base de datos MySQL se ejecutan directamente en la "
    "máquina virtual:")

add_styled_table(doc,
    ["Componente", "Host", "Puerto", "Protocolo"],
    [
        ["API REST (Spring Boot)", "52.201.91.206", "3000", "HTTP"],
        ["MySQL 8.x", "localhost", "3306", "TCP/JDBC"],
        ["App Móvil (desarrollo)", "Dispositivo/Emulador", "—", "—"],
        ["App Escritorio", "PC del administrador", "—", "—"],
        ["Terminal Python", "Kiosco/PC local", "—", "—"],
    ],
    col_widths=[4.5, 3.5, 2, 2.5]
)

add_styled_heading(doc, "11.2 Proceso de Despliegue Manual", level=2)

add_body(doc, "El despliegue actual sigue estos pasos:")

add_bullet(doc, "1. Compilar la API en local: ./gradlew build (genera un JAR ejecutable).")
add_bullet(doc, "2. Subir el JAR al servidor EC2 mediante SCP o wget desde el repositorio.")
add_bullet(doc, "3. Ejecutar los scripts SQL de inicialización de la base de datos.")
add_bullet(doc, "4. Arrancar el JAR: java -jar api-spring.jar (escucha en 0.0.0.0:3000).")
add_bullet(doc, "5. Configurar el Security Group de AWS para abrir el puerto 3000 (TCP).")

add_styled_heading(doc, "11.3 Dockerización (Preparado para Producción)", level=2)

add_body(doc,
    "El sistema está preparado para su dockerización con la siguiente "
    "configuración recomendada de docker-compose:")

add_code_block(doc,
"""# docker-compose.yml (configuración recomendada)
version: '3.8'

services:
  mysql:
    image: mysql:8.0
    container_name: reciapp-db
    environment:
      MYSQL_ROOT_PASSWORD: clase1234
      MYSQL_DATABASE: reciInventario_db
    ports:
      - "3306:3306"
    volumes:
      - mysql_data:/var/lib/mysql
      - ./sql/000_create_database.sql:/docker-entrypoint-initdb.d/01.sql
      - ./sql/001_imagen_longtext.sql:/docker-entrypoint-initdb.d/02.sql
    networks:
      - reciapp-net
    healthcheck:
      test: ["CMD", "mysqladmin", "ping", "-h", "localhost"]
      interval: 10s
      retries: 5

  api:
    build:
      context: ./api-spring
      dockerfile: Dockerfile
    container_name: reciapp-api
    ports:
      - "3000:3000"
    environment:
      SPRING_DATASOURCE_URL: jdbc:mysql://mysql:3306/reciInventario_db
      SPRING_DATASOURCE_USERNAME: root
      SPRING_DATASOURCE_PASSWORD: clase1234
    depends_on:
      mysql:
        condition: service_healthy
    networks:
      - reciapp-net

volumes:
  mysql_data:

networks:
  reciapp-net:
    driver: bridge""")

add_body(doc, "Dockerfile recomendado para la API:", bold=True)

add_code_block(doc,
"""# api-spring/Dockerfile
FROM eclipse-temurin:21-jre-alpine
WORKDIR /app
COPY build/libs/*.jar app.jar
EXPOSE 3000
ENTRYPOINT ["java", "-jar", "app.jar"]""")

add_styled_heading(doc, "11.4 Ventajas de la Dockerización", level=2)

add_bullet(doc, "Reproducibilidad: El entorno es idéntico en desarrollo y producción.")
add_bullet(doc, "Aislamiento: Cada servicio (API, MySQL) se ejecuta en su propio contenedor.")
add_bullet(doc, "Escalabilidad: Se pueden añadir réplicas de la API tras un balanceador de carga.")
add_bullet(doc, "Automatización: Un solo comando (docker compose up -d) levanta todo el sistema.")
add_bullet(doc, "Persistencia: Los datos de MySQL se almacenan en un volumen Docker que sobrevive a reinicios.")

add_styled_heading(doc, "11.5 Nota sobre HTTPS", level=2)

add_body(doc,
    "El sistema se comunica actualmente por HTTP. Se evaluó la implementación "
    "de HTTPS con certificados autofirmados, pero se descartó por las siguientes "
    "razones:",
    size=11)

add_bullet(doc, "Expo Go (React Native) no permite personalizar la configuración de seguridad de red de Android sin eyectar del flujo gestionado.")
add_bullet(doc, "Los certificados autofirmados requieren configuraciones adicionales en cada cliente (truststore en Java, urllib3 en Python).")
add_bullet(doc, "Para un entorno de producción real, se recomienda usar un certificado de una CA reconocida (como Let's Encrypt) detrás de un proxy Nginx.")

add_screenshot_placeholder(doc, 18, "Consola EC2 mostrando la API en ejecución (java -jar) o docker compose up")

doc.add_page_break()


# ═══════════════════════════════════════════════════════════════
# 12. CASOS DE USO
# ═══════════════════════════════════════════════════════════════

add_styled_heading(doc, "12. Casos de Uso", level=1)

add_styled_heading(doc, "12.1 Casos de Uso — Rol Cliente", level=2)

# CU-01
add_body(doc, "CU-01: Registro de Usuario", bold=True)
add_styled_table(doc,
    ["Campo", "Descripción"],
    [
        ["Actor", "Usuario no registrado"],
        ["Precondición", "El usuario no tiene cuenta en el sistema"],
        ["Flujo principal", "1. El usuario abre la app (móvil o escritorio).\n2. Pulsa 'Regístrate'.\n3. Introduce nombre de usuario y contraseña.\n4. El sistema valida que el nombre no exista.\n5. Se crea la cuenta con rol 'cliente' y emisiones a 0.\n6. Se genera un token JWT y se redirige a la pantalla principal."],
        ["Flujo alternativo", "4a. Si el nombre ya existe, se muestra un error."],
        ["Postcondición", "El usuario queda registrado y autenticado."],
    ],
    col_widths=[3, 12]
)

# CU-02
add_body(doc, "CU-02: Inicio de Sesión", bold=True)
add_styled_table(doc,
    ["Campo", "Descripción"],
    [
        ["Actor", "Usuario registrado (cliente o administrador)"],
        ["Precondición", "El usuario tiene una cuenta activa"],
        ["Flujo principal", "1. El usuario abre la app.\n2. Introduce nombre y contraseña.\n3. El sistema valida las credenciales con BCrypt.\n4. Se genera un token JWT.\n5. Se redirige a la pantalla principal (productos)."],
        ["Flujo alternativo", "3a. Credenciales incorrectas: se muestra un error."],
        ["Postcondición", "El usuario está autenticado con un token JWT válido."],
    ],
    col_widths=[3, 12]
)

# CU-03
add_body(doc, "CU-03: Consultar Catálogo de Productos", bold=True)
add_styled_table(doc,
    ["Campo", "Descripción"],
    [
        ["Actor", "Usuario autenticado (cliente/admin) o invitado"],
        ["Precondición", "El usuario ha iniciado sesión o está en modo invitado"],
        ["Flujo principal", "1. El usuario navega a la pestaña 'Productos'.\n2. Se muestra la lista completa de productos.\n3. Puede buscar por nombre o código de barras.\n4. Puede filtrar por material (PET, Vidrio, etc.)."],
        ["Postcondición", "El usuario visualiza los productos disponibles."],
    ],
    col_widths=[3, 12]
)

# CU-04
add_body(doc, "CU-04: Consultar Historial de Reciclaje", bold=True)
add_styled_table(doc,
    ["Campo", "Descripción"],
    [
        ["Actor", "Usuario autenticado"],
        ["Precondición", "El usuario ha iniciado sesión"],
        ["Flujo principal", "1. El usuario navega a la pestaña 'Historial'.\n2. Se muestra el total de CO₂ reducido.\n3. Se listan todas las transacciones de reciclaje con producto, fecha y emisiones."],
        ["Postcondición", "El usuario ve su impacto ambiental acumulado."],
    ],
    col_widths=[3, 12]
)

# CU-05
add_body(doc, "CU-05: Solicitar/Renovar TAP", bold=True)
add_styled_table(doc,
    ["Campo", "Descripción"],
    [
        ["Actor", "Usuario autenticado"],
        ["Precondición", "El usuario ha iniciado sesión"],
        ["Flujo principal", "1. El usuario accede a su perfil.\n2. Pulsa 'Ver TAP' o 'Solicitar TAP'.\n3. Se abre un modal mostrando el TAP actual o la opción de solicitarlo.\n4. Si pulsa 'Solicitar/Renovar', la API genera un número de 6 dígitos único.\n5. El nuevo TAP se muestra en el modal."],
        ["Postcondición", "El usuario tiene un TAP válido para usar en terminales."],
    ],
    col_widths=[3, 12]
)

# CU-06
add_body(doc, "CU-06: Modificar Datos Personales", bold=True)
add_styled_table(doc,
    ["Campo", "Descripción"],
    [
        ["Actor", "Usuario autenticado"],
        ["Precondición", "El usuario ha iniciado sesión"],
        ["Flujo principal", "1. El usuario accede a su perfil.\n2. Puede cambiar su nombre (requiere contraseña actual) o cambiar su contraseña (requiere actual + nueva + confirmación).\n3. La API valida la contraseña actual antes de aplicar cambios."],
        ["Flujo alternativo", "3a. Contraseña incorrecta: se muestra un error."],
        ["Postcondición", "Los datos del usuario se actualizan en la base de datos."],
    ],
    col_widths=[3, 12]
)

doc.add_page_break()

add_styled_heading(doc, "12.2 Casos de Uso — Rol Administrador", level=2)

# CU-07
add_body(doc, "CU-07: Gestión de Usuarios (CRUD)", bold=True)
add_styled_table(doc,
    ["Campo", "Descripción"],
    [
        ["Actor", "Administrador"],
        ["Precondición", "El admin ha iniciado sesión con rol 'administrador'"],
        ["Flujo principal", "1. El admin accede a la pestaña de administración (escritorio) o endpoints /api/admin/.\n2. Puede listar todos los usuarios del sistema.\n3. Puede crear nuevos usuarios con nombre, contraseña y rol.\n4. Puede modificar nombre, contraseña, rol y emisiones de cualquier usuario.\n5. Puede eliminar usuarios (CASCADE elimina su historial)."],
        ["Postcondición", "Los usuarios del sistema se modifican según la acción realizada."],
    ],
    col_widths=[3, 12]
)

# CU-08
add_body(doc, "CU-08: Gestión de Productos (CRUD)", bold=True)
add_styled_table(doc,
    ["Campo", "Descripción"],
    [
        ["Actor", "Administrador"],
        ["Precondición", "El admin ha iniciado sesión con rol 'administrador'"],
        ["Flujo principal", "1. El admin accede a la gestión de productos.\n2. Puede listar todos los productos con sus datos.\n3. Puede crear nuevos productos con tipo, código de barras, nombre, material, emisiones e imagen.\n4. Puede modificar cualquier campo de un producto existente.\n5. Puede eliminar productos del catálogo."],
        ["Postcondición", "El catálogo de productos se actualiza."],
    ],
    col_widths=[3, 12]
)

# CU-09
add_body(doc, "CU-09: Gestión de Transacciones (CRUD)", bold=True)
add_styled_table(doc,
    ["Campo", "Descripción"],
    [
        ["Actor", "Administrador"],
        ["Precondición", "El admin ha iniciado sesión con rol 'administrador'"],
        ["Flujo principal", "1. El admin accede a la gestión de transacciones.\n2. Puede listar todas las transacciones de reciclaje de todos los usuarios.\n3. Puede crear transacciones manualmente (asignar reciclaje a un usuario).\n4. Puede eliminar transacciones incorrectas."],
        ["Postcondición", "El historial de reciclaje se modifica según la acción."],
    ],
    col_widths=[3, 12]
)

add_styled_heading(doc, "12.3 Casos de Uso — Terminal de Reciclaje", level=2)

# CU-10
add_body(doc, "CU-10: Registrar Reciclaje en Terminal", bold=True)
add_styled_table(doc,
    ["Campo", "Descripción"],
    [
        ["Actor", "Usuario + Operador del terminal"],
        ["Precondición", "El terminal está conectado a la API y el operador autenticado"],
        ["Flujo principal", "1. El usuario coloca el producto en el punto de reciclaje.\n2. Se escanea el código de barras del producto.\n3. El terminal busca el producto en la API.\n4. El usuario introduce su TAP (6 dígitos).\n5. El terminal verifica el TAP en la API.\n6. Se muestra un resumen y se pide confirmación.\n7. Se registra el reciclaje y se muestra la confirmación con emisiones acumuladas."],
        ["Flujo alternativo", "3a. Producto no encontrado: se muestra error.\n5a. TAP no válido: se muestra error."],
        ["Postcondición", "El reciclaje queda registrado y las emisiones del usuario se actualizan."],
    ],
    col_widths=[3, 12]
)

add_styled_heading(doc, "12.4 Casos de Uso — Modo Invitado (App Móvil)", level=2)

# CU-11
add_body(doc, "CU-11: Explorar en Modo Offline/Invitado", bold=True)
add_styled_table(doc,
    ["Campo", "Descripción"],
    [
        ["Actor", "Usuario sin conexión o sin cuenta"],
        ["Precondición", "La app no puede conectar a la API o el usuario elige modo offline"],
        ["Flujo principal", "1. Se muestra el modal de conexión con opciones: Recargar o Modo Offline.\n2. El usuario elige Modo Offline.\n3. Se cargan datos de ejemplo desde el almacenamiento local.\n4. El usuario puede explorar productos de ejemplo.\n5. El historial y perfil muestran banners indicando el modo invitado."],
        ["Postcondición", "El usuario explora la app con funcionalidad limitada."],
    ],
    col_widths=[3, 12]
)

doc.add_page_break()


# ═══════════════════════════════════════════════════════════════
# 13. CAPTURAS DE PANTALLA
# ═══════════════════════════════════════════════════════════════

add_styled_heading(doc, "13. Capturas de Pantalla", level=1)

add_body(doc,
    "A continuación se listan todas las capturas de pantalla que deben "
    "incluirse en este documento. Cada recuadro indica la captura necesaria. "
    "Sustituye cada recuadro por la captura correspondiente.")

add_body(doc, "")

# Resumen de todas las capturas que ya hemos incluido
capturas_resumen = [
    ("Arquitectura", "Diagrama general de la arquitectura del sistema"),
    ("Autenticación", "Flujo de autenticación JWT (diagrama de secuencia)"),
    ("Base de datos", "Modelo Entidad-Relación desde MySQL Workbench"),
    ("API REST", "Petición en Postman con respuesta JSON"),
    ("App Móvil — Login", "Pantalla de login de la app móvil"),
    ("App Móvil — Productos", "Pantalla de productos con filtro de materiales"),
    ("App Móvil — Historial", "Pantalla de historial con total CO₂"),
    ("App Móvil — Perfil", "Pantalla de perfil con estadísticas y TAP"),
    ("App Móvil — Offline", "Modal de conexión offline"),
    ("JavaFX — Login", "Pantalla de login de la app de escritorio"),
    ("JavaFX — Info Personal", "Dashboard - pestaña información personal"),
    ("JavaFX — Admin Usuarios", "Dashboard - administración de usuarios"),
    ("JavaFX — Admin Productos", "Dashboard - administración de productos"),
    ("JavaFX — Gráfico", "Gráfico de emisiones por usuario"),
    ("JavaFX — Diálogo", "Ventana de diálogo crear/modificar producto"),
    ("Terminal — Banner", "Terminal: banner de bienvenida y autenticación"),
    ("Terminal — Reciclaje", "Terminal: flujo completo de reciclaje"),
    ("Despliegue", "Consola EC2 con la API en ejecución"),
]

add_styled_table(doc,
    ["#", "Sección", "Descripción"],
    [[str(i+1), cap[0], cap[1]] for i, cap in enumerate(capturas_resumen)],
    col_widths=[1.2, 4, 9]
)

# ═══════════════════════════════════════════════════════════════
# Guardar documento
# ═══════════════════════════════════════════════════════════════

output_path = os.path.join(os.path.dirname(__file__), "ReciApp_Documentacion_Proyecto.docx")
doc.save(output_path)
print(f"Documento generado correctamente: {output_path}")
print(f"   Capturas de pantalla: {screenshot_counter}")
